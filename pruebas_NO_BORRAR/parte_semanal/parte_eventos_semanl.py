import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pypdf import PdfReader
import re
import pandas as pd
from docx import Document
import openpyxl
from openpyxl.styles import Font, Alignment
import traceback
from datetime import datetime
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import threading
import queue
import os

# ==========================================
# CONFIGURATION & CONSTANTS
# ==========================================


class Config:
    MESES_ESPANOL = {
        "enero": 1,
        "ene": 1,
        "febrero": 2,
        "feb": 2,
        "marzo": 3,
        "mar": 3,
        "abril": 4,
        "abr": 4,
        "mayo": 5,
        "may": 5,
        "junio": 6,
        "jun": 6,
        "julio": 7,
        "jul": 7,
        "agosto": 8,
        "ago": 8,
        "septiembre": 9,
        "sep": 9,
        "setiembre": 9,
        "octubre": 10,
        "oct": 10,
        "noviembre": 11,
        "nov": 11,
        "diciembre": 12,
        "dic": 12,
    }

    DIAS_SEMANA = {
        0: "LUNES",
        1: "MARTES",
        2: "MIÉRCOLES",
        3: "JUEVES",
        4: "VIERNES",
        5: "SÁBADO",
        6: "DOMINGO",
    }

    MAPA_NUMEROS = {
        "un": 1,
        "uno": 1,
        "una": 1,
        "dos": 2,
        "tres": 3,
        "cuatro": 4,
        "cinco": 5,
        "seis": 6,
        "siete": 7,
        "ocho": 8,
        "nueve": 9,
        "diez": 10,
    }

    # Compilation of Regex Patterns for efficiency
    RE_FECHA_TEXTO = re.compile(
        r"(\d{1,2})\s+de\s+([A-záéíóúñ]+)(?:\s+de(?:l)?\s+(\d{2,4}))?", re.IGNORECASE
    )
    RE_FECHA_TEXTO_SIMPLE = re.compile(
        r"(\d{1,2})\s+([A-záéíóúñ]+)(?:\s+de(?:l)?\s+(\d{2,4}))?", re.IGNORECASE
    )
    RE_FECHA_NUMERICA = re.compile(r"(\d{1,2})[\/\-](\d{1,2})(?:[\/\-](\d{2,4}))?")
    RE_FECHA_ISO = re.compile(r"(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})")

    # Data Extraction Patterns
    RE_FECHA_DOC = re.compile(r"Fecha\s*:\s*(.+?)(?:\s+Hora\b|\n|$)", re.IGNORECASE)
    RE_HORA_RESENA = re.compile(
        r"(breve\s+reseña|reseña)\s*[:\-]?\s*(.*?)\b(?:siendo\s+las|a\s+las)\s*(\d{1,2}:\d{2})",
        re.IGNORECASE | re.DOTALL,
    )
    RE_HORA_GENERAL = re.compile(
        r"(hora|horario)\s*[:\-]?\s*(\d{1,2}:\d{2})", re.IGNORECASE
    )
    RE_CAUSA = re.compile(
        r"^\s*(?:causa|hecho)\b\s*[\s*:\-–—]?\s*(.+?)\s*$", re.IGNORECASE | re.MULTILINE
    )
    RE_JEFE = re.compile(r"j\s*e\s*f\s*e\s*de\s*servicio:\s*([^\n]+)", re.IGNORECASE)
    RE_OFICIAL = re.compile(
        r"(?:Oficial\s*+de\s*+Serv[ií]cio:\s*|Oficial\s*+de\s*+Serv[ií]cio\s*)([^\n]+)",
        re.IGNORECASE,
    )
    RE_OPERADOR = re.compile(
        r"(?:Operador\s*de\s*C[aá]mara|Operador\s*C[aá]mara|Op(?:erador)?\s*C[aá]mara|Aux(?:iliar)?\s+Operador\s+de\s+C[aá]mara|Aux\s*)\s*:?\s*([^\n]+)",
        re.IGNORECASE,
    )
    RE_SAE = re.compile(
        r"(?:SAE\s*Nro.\s*|SAE:\s*°\s*|SAE\s*nº\s*|sae\s*|suceso\s*:\s*|suceso\s*|cad\s*:\s*|SAECAD\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*|sae\s*+nro.|carta\s*:|carta\s*n\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*|SAE\s*N.\s*º\s*|N°\s*|N°\s*:\s*)\s*(\d{8})",
        re.IGNORECASE,
    )
    RE_RESULTADO = re.compile(r"\bresultado\b\s*[:\-–—]?\s*([^\n]+)", re.IGNORECASE)
    RE_RESENA_BLOQUE = re.compile(
        r"(?:rese[nñ]a|rese[nñ]a)\s*:?\s*(.*?)Resultado\s*:", re.IGNORECASE | re.DOTALL
    )


# ==========================================
# PROCESSING LOGIC
# ==========================================


class PDFProcessor:
    @staticmethod
    def formatear_fecha(fecha_texto):
        if not fecha_texto or fecha_texto == "S/D":
            return "S/D"

        fecha_limpia = fecha_texto.strip()
        anio_actual = datetime.now().year

        try:
            # 1. Text format
            m = Config.RE_FECHA_TEXTO.search(fecha_limpia)
            if not m:
                m = Config.RE_FECHA_TEXTO_SIMPLE.search(fecha_limpia)

            if m:
                dia = int(m.group(1))
                mes_texto = m.group(2).lower().replace(".", "")
                mes = Config.MESES_ESPANOL.get(mes_texto)
                if mes:
                    anio = int(m.group(3)) if m.group(3) else anio_actual
                    if anio < 100:
                        anio += 2000
                    return f"{dia:02d}/{mes:02d}/{anio}"

            # 2. Numeric format
            m = Config.RE_FECHA_NUMERICA.search(fecha_limpia)
            if m:
                dia = int(m.group(1))
                mes = int(m.group(2))
                anio = int(m.group(3)) if m.group(3) else anio_actual
                if anio < 100:
                    anio += 2000
                return f"{dia:02d}/{mes:02d}/{anio}"

            # 3. ISO format
            m = Config.RE_FECHA_ISO.search(fecha_limpia)
            if m:
                return f"{int(m.group(3)):02d}/{int(m.group(2)):02d}/{int(m.group(1))}"

            # 4. Standard formats with datetime
            formatos = ["%d/%m/%Y", "%d/%m/%y", "%Y-%m-%d", "%m/%d/%Y", "%d-%m-%Y"]
            for fmt in formatos:
                try:
                    return datetime.strptime(fecha_limpia, fmt).strftime("%d/%m/%Y")
                except ValueError:
                    continue

            return fecha_limpia
        except Exception as e:
            print(f"Error parsing date '{fecha_texto}': {e}")
            return fecha_texto

    @staticmethod
    def obtener_dia_semana(fecha_str):
        if not fecha_str or fecha_str == "S/D":
            return "S/D"
        try:
            fecha_obj = None
            for fmt in ["%d/%m/%Y", "%d/%m/%y"]:
                try:
                    fecha_obj = datetime.strptime(fecha_str, fmt)
                    break
                except ValueError:
                    continue

            if not fecha_obj:
                # Try adding current year if missing
                if re.match(r"\d{1,2}/\d{1,2}$", fecha_str):
                    fecha_obj = datetime.strptime(
                        f"{fecha_str}/{datetime.now().year}", "%d/%m/%Y"
                    )

            if fecha_obj:
                return Config.DIAS_SEMANA[fecha_obj.weekday()]
        except Exception:
            pass
        return "S/D"

    @staticmethod
    def extraer_numero_resultado(texto):
        t = texto.lower()
        # Look for "number (" pattern
        m = re.search(r"\b(\d+)\s*\(", t)
        if m:
            return int(m.group(1))
        # Look for standalone number
        m = re.search(r"\b(\d+)\b", t)
        if m:
            return int(m.group(1))
        # Look for "(number)"
        m = re.search(r"\((\d+)\)", t)
        if m:
            return int(m.group(1))
        # Look for text numbers
        for palabra, valor in Config.MAPA_NUMEROS.items():
            if re.search(rf"\b{palabra}\b", t):
                return valor
        return ""

    @classmethod
    def procesar_pdf(cls, ruta):
        try:
            reader = PdfReader(ruta)
            texto = "\n".join(
                [p.extract_text() for p in reader.pages if p.extract_text()]
            )
        except Exception as e:
            raise Exception(f"No se pudo leer el PDF: {e}")

        # Fecha
        fecha_match = Config.RE_FECHA_DOC.search(texto)
        fecha_texto = fecha_match.group(1).strip() if fecha_match else "S/D"
        # Fallback date search if specific field not found
        if fecha_texto == "S/D":
            fallback = re.search(r"(\d{1,2}\s+de\s+[A-Za-z]+\s+de\s+\d{4})", texto)
            if fallback:
                fecha_texto = fallback.group(1)

        fecha_fmt = cls.formatear_fecha(fecha_texto)
        dia_semana = cls.obtener_dia_semana(fecha_fmt)

        # Hora
        hora_match = Config.RE_HORA_RESENA.search(texto)
        hora = hora_match.group(3) if hora_match else "S/D"
        if hora == "S/D":
            h_gen = Config.RE_HORA_GENERAL.search(texto)
            if h_gen:
                hora = h_gen.group(2)

        # Campos simples
        causa = Config.RE_CAUSA.search(texto)
        causa = causa.group(1).strip() if causa else "S/D"

        jefe = Config.RE_JEFE.search(texto)
        jefe = jefe.group(1).strip() if jefe else "S/D"

        oficial = Config.RE_OFICIAL.search(texto)
        oficial = oficial.group(1).strip() if oficial else "S/D"

        operador = Config.RE_OPERADOR.search(texto)
        op_val = operador.group(1).strip() if operador else "S/D"
        # Sanitize operador value
        if re.match(
            r"^(breve|rese[nñ]a|resultado|fecha|hora)\b", op_val, re.IGNORECASE
        ):
            op_val = "S/D"

        sae = Config.RE_SAE.search(texto)
        sae_val = int(sae.group(1).strip()) if sae else "S/D"

        resultado = Config.RE_RESULTADO.search(texto)
        res_texto = resultado.group(1).strip() if resultado else "S/D"

        resena = Config.RE_RESENA_BLOQUE.search(texto)
        resena_val = re.sub(r"\s+", " ", resena.group(1)).strip() if resena else "S/D"

        # Conteo de personas
        det = (
            cls.extraer_numero_resultado(res_texto)
            if re.search(r"\bdeten", res_texto, re.IGNORECASE)
            else ""
        )
        cont = (
            cls.extraer_numero_resultado(res_texto)
            if re.search(r"\bcontra", res_texto, re.IGNORECASE)
            else ""
        )
        imp = (
            cls.extraer_numero_resultado(res_texto)
            if re.search(r"\bimputad", res_texto, re.IGNORECASE)
            else ""
        )

        res_final = "" if (det or cont or imp) else res_texto

        return {
            "RUTA": ruta,
            "FECHA": fecha_fmt,
            "DIA_SEMANA": dia_semana,
            "CAUSA": causa,
            "SAE": sae_val,
            "HORA": hora,
            "OPERADOR DE CAMARA": op_val,
            "OFICIAL DE SERVICIO": oficial,
            "JEFE DE SERVICIO": jefe,
            "RESULTADO_TEXTO": res_final,
            "RESEÑA": resena_val,
            "DETENIDO": det,
            "CONTRAVENTOR": cont,
            "IMPUTADO": imp,
        }


# ==========================================
# EXPORT LOGIC
# ==========================================


class DataExporter:
    @staticmethod
    def agregar_bordes_tabla(tabla):
        def set_border(celda):
            tc = celda._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")
                tcBorders.append(border)
            tcPr.append(tcBorders)

        for row in tabla.rows:
            for cell in row.cells:
                set_border(cell)

    @staticmethod
    def guardar_excel(datos_lista, filename="contenido_archivos.xlsx"):
        # Map internal keys to Excel headers
        excel_data = []
        for d in datos_lista:
            excel_data.append(
                {
                    "RUTA": d["RUTA"],
                    "FECHA": d["FECHA"],
                    "TIPO DE EVENTO": "",
                    "CAUSA": d["CAUSA"],
                    "SAE": d["SAE"],
                    "HORA": d["HORA"],
                    "OPERADOR DE CAMARA": d["OPERADOR DE CAMARA"],
                    "OFICIAL DE SERVICIO": d["OFICIAL DE SERVICIO"],
                    "CONTRAVENTOR": d["CONTRAVENTOR"],
                    "DETENIDO": d["DETENIDO"],
                    "IMPUTADO": d["IMPUTADO"],
                    "OBSERVACIONES": "",
                    "GRUPO QUE TRABAJO ESE DÍA": "",
                    "JEFE DE SERVICIO": d["JEFE DE SERVICIO"],
                    "RESULTADO": d["RESULTADO_TEXTO"],
                }
            )

        df_nuevo = pd.DataFrame(excel_data)

        try:
            if os.path.exists(filename):
                df_existente = pd.read_excel(filename)
                df_final = pd.concat([df_existente, df_nuevo], ignore_index=True)
            else:
                df_final = df_nuevo

            df_final.to_excel(filename, index=False)
            DataExporter.formatear_excel(filename)
        except PermissionError:
            raise PermissionError(
                f"No se puede guardar '{filename}'. Cierra el archivo si está abierto."
            )
        except Exception as e:
            raise Exception(f"Error guardando Excel: {e}")

    @staticmethod
    def formatear_excel(filename):
        wb = openpyxl.load_workbook(filename)
        ws = wb.active

        # Cleanup regex
        palabras_limpiar = [
            "personal",
            "contratado",
            "auxiliar",
            "pers",
            "contr",
            "aux",
            "cont",
            "oficial",
            "of",
            "operador",
            "oper",
            "lp",
            "primero",
            "mayor",
            "inspector",
            "subrio",
            "subcomisario",
        ]
        regex_limpiar = re.compile(
            r"\b(" + "|".join(re.escape(p) for p in palabras_limpiar) + r")\.?\b",
            re.IGNORECASE,
        )

        cols_limpiar = ["G", "H", "N"]  # 1-based indices: 7, 8, 14
        cols_center = ["B", "E", "F", "I", "J", "K"]  # 2, 5, 6, 9, 10, 11

        # Helper to get column letter -> index 1-based (not strictly needed with openpyxl strings, but good for robust logic if needed)

        for row in ws.iter_rows(min_row=2):
            # Cleaning columns
            for col_letter in cols_limpiar:
                try:
                    cell = ws[f"{col_letter}{row[0].row}"]
                    if cell.value:
                        val = str(cell.value)
                        val = regex_limpiar.sub("", val)
                        val = re.sub(r"[.\-:;–]", " ", val)
                        val = re.sub(r"\d+", " ", val)
                        val = re.sub(r"\s+", " ", val).strip()
                        cell.value = val if val else "S/D"
                except:
                    pass

            # Formatting Date
            cell_fecha = ws[f"B{row[0].row}"]
            if cell_fecha.value and str(cell_fecha.value) != "S/D":
                try:
                    # Try parsing date string to object for Excel
                    dt = datetime.strptime(str(cell_fecha.value), "%d/%m/%Y")
                    cell_fecha.value = dt
                    cell_fecha.number_format = "DD/MM/YYYY"
                except:
                    pass

            # Hyperlinks
            cell_ruta = ws[f"A{row[0].row}"]
            if cell_ruta.value and str(cell_ruta.value).lower().endswith(".pdf"):
                cell_ruta.hyperlink = cell_ruta.value
                cell_ruta.font = Font(color="0000FF", underline="single")

            # Allignment
            for col in cols_center:
                ws[f"{col}{row[0].row}"].alignment = Alignment(horizontal="center")

        wb.save(filename)

    @staticmethod
    def crear_word(datos_agrupados, filename="Parte_Semanal_Eventos.docx"):
        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # --- Title ---
        p = doc.add_paragraph()
        run = p.add_run("División CENTROS DE MONITOREO URBANO")
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # --- Period ---
        fechas = sorted([f for f in datos_agrupados.keys()])
        texto_periodo = "PARTE SEMANAL"
        if fechas:
            texto_periodo += f" DEL DÍA {fechas[0]} al DÍA {fechas[-1]}"

        p = doc.add_paragraph()
        run = p.add_run(texto_periodo)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # --- Table ---
        tabla = doc.add_table(rows=1, cols=5)  # Start with header
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.style = "Table Grid"

        headers = [
            "MES DE AGOSTO/DÍA",
            "EVENTOS DEL PERÍODO",
            "EVENTOS POSITIVOS",
            "DETENIDOS",
            "EVENTOS RELEVANTES",
        ]
        hdr_cells = tabla.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
            if i == 0:
                hdr_cells[i].width = Inches(1.5)

        # Table Rows
        dias_semana_orden = [
            "LUNES",
            "MARTES",
            "MIÉRCOLES",
            "JUEVES",
            "VIERNES",
            "SÁBADO",
            "DOMINGO",
        ]
        stats_totales = {"det": 0, "resenas": 0, "eventos": 0}

        # Pre-process data to map by weekday
        mapa_sem = {dia: {"det": 0, "resenas": 0} for dia in dias_semana_orden}

        for fecha, data in datos_agrupados.items():
            dia = data["dia_semana"]
            if dia in mapa_sem:
                mapa_sem[dia]["det"] += data["total_personas"]
                mapa_sem[dia]["resenas"] += len(data["reseñas"])
                stats_totales["det"] += data["total_personas"]
                stats_totales["resenas"] += len(data["reseñas"])
                stats_totales["eventos"] += len(data["eventos"])

        for dia in dias_semana_orden:
            row = tabla.add_row().cells
            # Cell 1: Day
            row[0].text = dia
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[0].paragraphs[0].runs[0].bold = True
            row[0].paragraphs[0].runs[0].font.size = Pt(9)

            # Empty cells
            row[1].text = ""
            row[2].text = ""

            # Counts
            d = mapa_sem[dia]
            row[3].text = str(d["det"]) if d["det"] > 0 else ""
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row[4].text = str(d["resenas"]) if d["resenas"] > 0 else ""
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Total Row
        row = tabla.add_row().cells
        row[0].text = "TOTAL"
        row[0].paragraphs[0].runs[0].bold = True
        row[3].text = str(stats_totales["det"]) if stats_totales["det"] > 0 else ""
        row[3].paragraphs[0].runs[0].bold = True
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[4].text = (
            str(stats_totales["resenas"]) if stats_totales["resenas"] > 0 else ""
        )
        row[4].paragraphs[0].runs[0].bold = True
        row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        DataExporter.agregar_bordes_tabla(tabla)
        for row in tabla.rows:
            row.height = Inches(0.4)

        doc.add_page_break()

        # --- Detailed Reviews ---
        p = doc.add_paragraph()
        run = p.add_run("RESEÑAS COMPLETAS POR DÍA")
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for fecha in sorted(datos_agrupados.keys()):
            data = datos_agrupados[fecha]
            if not data["reseñas"]:
                continue

            p = doc.add_paragraph()
            run = p.add_run(f"{data['dia_semana']} - {fecha}")
            run.bold = True
            run.font.size = Pt(12)

            '''p_stats = doc.add_paragraph()
            txt_stats = f"Total reseñas: {len(data['reseñas'])} | Total personas: {data['total_personas']} (D:{data['detenidos']}, C:{data['contraventores']}, I:{data['imputados']})"
            p_stats.add_run(txt_stats).italic = True'''

            for idx, resena in enumerate(data["reseñas"], 1):
                p_r = doc.add_paragraph()
                p_r.add_run(f"{idx}. ").bold = True
                p_r.add_run(resena)

            p_sep = doc.add_paragraph()
            p_sep.add_run("-" * 50).italic = True
            p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            doc.save(filename)
        except PermissionError:
            raise PermissionError(
                f"No se puede guardar '{filename}'. Cierra el archivo si está abierto."
            )


# ==========================================
# APP UI
# ==========================================


class ParteSemanalApp:
    def __init__(self, root):
        self.root = root
        self.rutas_pdf = []

        # UI Setup
        self.root.title("Extracción de datos .PDF - Parte Semanal")
        self.root.config(bg="#822121")
        self.root.geometry("840x450")

        main_frame = tk.Frame(root, bg="#822121")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Status Label
        self.lbl_status = tk.Label(
            main_frame,
            text="Seleccione archivos PDF para comenzar",
            bg="#1a1a1a",
            fg="#00FF00",
            font=("Helvetica", 11),
            relief="sunken",
            height=8,
            wraplength=800,
            justify="left",
        )
        self.lbl_status.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        # Progress Bar
        self.progress = ttk.Progressbar(main_frame, mode="indeterminate")
        self.progress.pack(fill=tk.X, pady=(0, 10))

        # Buttons Frame
        btn_frame = tk.Frame(main_frame, bg="#822121")
        btn_frame.pack(fill=tk.X)

        self.btn_select = tk.Button(
            btn_frame,
            text="📁 Seleccionar PDFs",
            command=self.seleccionar_archivos,
            bg="#33AAFF",
            fg="white",
            font=("Helvetica", 10, "bold"),
            height=2,
        )
        self.btn_select.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)

        self.btn_run = tk.Button(
            btn_frame,
            text="▶ Ejecutar",
            command=self.iniciar_procesamiento,
            bg="#33FF41",
            fg="white",
            font=("Helvetica", 10, "bold"),
            height=2,
        )
        self.btn_run.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)

        self.btn_close = tk.Button(
            btn_frame,
            text="❌ Cerrar",
            command=root.quit,
            bg="#FF3333",
            fg="white",
            font=("Helvetica", 10, "bold"),
            height=2,
        )
        self.btn_close.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)

    def log(self, mensaje):
        def _update():
            current_text = self.lbl_status.cget("text")
            # Keep last few lines
            lines = current_text.split("\n")[-10:]
            lines.append(mensaje)
            self.lbl_status.config(text="\n".join(lines))
            print(mensaje)
        self.root.after(0, _update)

    def seleccionar_archivos(self):
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")],
        )
        if archivos:
            self.rutas_pdf = list(archivos)
            self.lbl_status.config(
                text=f"Archivos seleccionados:\n"
                + "\n".join([os.path.basename(f) for f in archivos[:5]])
            )
            if len(archivos) > 5:
                self.log(f"... y {len(archivos) - 5} más.")

    def iniciar_procesamiento(self):
        if not self.rutas_pdf:
            messagebox.showwarning("Aviso", "Seleccione archivos primero.")
            return

        self.btn_run.config(state="disabled")
        self.btn_select.config(state="disabled")
        self.progress.start(10)
        self.log("Iniciando procesamiento...")

        # Run in thread
        threading.Thread(target=self.procesar_thread, daemon=True).start()

    def procesar_thread(self):
        try:
            resultados = []
            datos_por_dia = {}
            error_msg = None

            for ruta in self.rutas_pdf:
                try:
                    data = PDFProcessor.procesar_pdf(ruta)
                    resultados.append(data)

                    # Grouping logic
                    fecha = data["FECHA"]
                    dia = data["DIA_SEMANA"]

                    if fecha != "S/D" and dia != "S/D":
                        if fecha not in datos_por_dia:
                            datos_por_dia[fecha] = {
                                "dia_semana": dia,
                                "eventos": [],
                                "reseñas": [],
                                "detenidos": 0,
                                "contraventores": 0,
                                "imputados": 0,
                                "total_personas": 0,
                            }

                        # Add stats
                        det = (
                            data["DETENIDO"] if isinstance(data["DETENIDO"], int) else 0
                        )
                        cont = (
                            data["CONTRAVENTOR"]
                            if isinstance(data["CONTRAVENTOR"], int)
                            else 0
                        )
                        imp = (
                            data["IMPUTADO"] if isinstance(data["IMPUTADO"], int) else 0
                        )
                        total = det + cont + imp

                        datos_por_dia[fecha]["detenidos"] += det
                        datos_por_dia[fecha]["contraventores"] += cont
                        datos_por_dia[fecha]["imputados"] += imp
                        datos_por_dia[fecha]["total_personas"] += total
                        datos_por_dia[fecha]["eventos"].append(data)

                        if data["RESEÑA"] and data["RESEÑA"] != "S/D":
                            datos_por_dia[fecha]["reseñas"].append(data["RESEÑA"])

                except Exception as e:
                    self.log(f"Error procesando {os.path.basename(ruta)}: {e}")

            # Exporting
            if resultados:
                self.log("Guardando Excel...")
                DataExporter.guardar_excel(resultados)

                self.log("Generando Word...")
                DataExporter.crear_word(datos_por_dia)
            else:
                error_msg = "No se extrajeron datos de los PDFs."

        except Exception as e:
            error_msg = str(e)
            traceback.print_exc()
        finally:
            self.root.after(0, self.finalizar_flujo, error_msg)

    def finalizar_flujo(self, error):
        self.progress.stop()
        self.btn_run.config(state="normal")
        self.btn_select.config(state="normal")

        if error:
            messagebox.showerror("Error", f"Ocurrió un error:\n{error}")
            self.log(f"ERROR: {error}")
        else:
            messagebox.showinfo("Éxito", "Proceso completado correctamente.")
            self.log("Proceso finalizado. Archivos guardados.")


if __name__ == "__main__":
    root = tk.Tk()
    app = ParteSemanalApp(root)
    root.mainloop()
