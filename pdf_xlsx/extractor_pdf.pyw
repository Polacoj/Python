import tkinter as tk
import re
import pandas as pd
import openpyxl
import traceback
from pypdf import PdfReader
from tkinter import filedialog, messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Alignment
from openpyxl.worksheet.hyperlink import Hyperlink
"""
Extractor de datos desde archivos PDF y conversión a Excel y Word"""
# Variable global para almacenar rutas
rutas_global = []

def formatear_fecha(fecha_texto):
    """
    Convierte diferentes formatos de fecha y devuelve solo 'dd/mm'.
    Si no se puede parsear, devuelve el texto original.
    """
    if not fecha_texto or fecha_texto == "S/D":
        return "S/D"
    try:
        fecha_limpia = fecha_texto.strip()
        meses_espanol = {
            'enero': 1, 'ene': 1,
            'febrero': 2, 'feb': 2,
            'marzo': 3, 'mar': 3,
            'abril': 4, 'abr': 4,
            'mayo': 5, 'may': 5,
            'junio': 6, 'jun': 6,
            'julio': 7, 'jul': 7,
            'agosto': 8, 'ago': 8,
            'septiembre': 9, 'sep': 9, 'setiembre': 9,
            'octubre': 10, 'oct': 10,
            'noviembre': 11, 'nov': 11,
            'diciembre': 12, 'dic': 12
        }

        # 1) Formato con mes en texto: "23 de octubre", "23 octubre 2025", "23 de oct 2025"
        m = re.search(r'(\d{1,2})\s+de\s+([A-záéíóúñ]+)', fecha_limpia, re.IGNORECASE)
        if not m:
            m = re.search(r'(\d{1,2})\s+([A-záéíóúñ]+)', fecha_limpia, re.IGNORECASE)
        if m:
            dia = int(m.group(1))
            mes_texto = m.group(2).lower()
            mes_texto = mes_texto.replace('.', '')
            mes = meses_espanol.get(mes_texto)
            if mes:
                return f"{dia:02d}/{mes:02d}"

        # 2) Formato numérico dd/mm[/aaaa] o dd-mm[-aaaa]
        m = re.search(r'(\d{1,2})[\/\-](\d{1,2})(?:[\/\-]\d{2,4})?', fecha_limpia)
        if m:
            dia = int(m.group(1))
            mes = int(m.group(2))
            return f"{dia:02d}/{mes:02d}"

        # 3) Formato ISO aaaa-mm-dd -> extraer día/mes
        m = re.search(r'(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})', fecha_limpia)
        if m:
            dia = int(m.group(3))
            mes = int(m.group(2))
            return f"{dia:02d}/{mes:02d}"

        # 4) Intentar parseo con datetime para formatos comunes y extraer día/mes
        for fmt in ("%d/%m/%Y", "%d/%m", "%Y-%m-%d", "%m-%d"):
            try:
                fecha_obj = datetime.strptime(fecha_limpia, fmt)
                return fecha_obj.strftime("%d/%m")
            except Exception:
                continue

        # Si no se pudo convertir, devolver el texto original
        return fecha_limpia
    except Exception as e:
        print(f"Error al formatear fecha '{fecha_texto}': {e}")
        return fecha_texto

def seleccionar_archivos():
    """Seleccionar archivos PDF usando diálogo de archivos"""
    """AI is creating summary for seleccionar_archivos
    """
    try:
        global rutas_global
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")]
        )
        
        if archivos:
            rutas_global = list(archivos)
            archivos_texto = "Archivos seleccionados:\n" + "\n".join([f.split("/")[-1] for f in archivos])
            ventana_pegado.config(text=archivos_texto)
            print(f"Archivos seleccionados: {archivos}")
        else:
            ventana_pegado.config(text="No se seleccionaron archivos")
    except Exception as e:
        print(f"Error en selección de archivos: {e}")
        traceback.print_exc()
        messagebox.showerror("Error", f"Error al seleccionar archivos: {str(e)}")

def extraer_convertir():
    """
    Docstring para extraer_convertir
    
    Extraer datos de archivos PDF y guardarlos en Excel y Word"""
    try:
        if not rutas_global:
            messagebox.showwarning("Advertencia", "Por favor, selecciona archivos PDF primero.")
            return
            
        datos, datos_2 = [], []
        error = ""
        
        # Procesar cada PDF por separado en lugar de concatenar y split
        for ruta in rutas_global:
            try:
                documento = PdfReader(ruta)
                texto = ""
                for pagina in documento.pages:
                    texto_pagina = pagina.extract_text()
                    if texto_pagina:
                        texto += texto_pagina + "\n"
            except Exception as e:
                error += f"Error al leer {ruta}: {str(e)}\n"
                texto = ""
            
            # Extracciones por archivo (misma lógica que antes, aplicada a 'texto' y 'ruta')
            fecha = re.search(r'Fecha\s*:\s*('
                    r'\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}'
                    r'\d{1,2}\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}'
                    r'|\d{1,2}[/-]\d{1,2}[/-]\d{4}'
                    r'|\d{4}[/-]\d{1,2}[/-]\d{1,2}'
                    r')', 
                    texto,
                    re.IGNORECASE)
            # 1) Hora dentro del bloque de reseña
            hora_resena = re.search(
                        r'(breve\s+reseña|reseña)\s*[:\-]?\s*(.*?)\b(?:siendo\s+las|a\s+las)\s*(\d{1,2}:\d{2})',
                        texto,
                        re.IGNORECASE | re.DOTALL
                        )

            # 2) Hora general
            hora_general = re.search(
                            r'(hora|horario)\s*[:\-]?\s*(\d{1,2}:\d{2})',
                            texto,
                            re.IGNORECASE
                        )

            # 3) Selección final
            if hora_resena:
                hora_final = hora_resena.group(3)
            elif hora_general:
                hora_final = hora_general.group(2)
            else:
                hora_final = "S/D"

            causa = re.search(r'^\s*(?:causa|hecho)\b\s*[\s*:\-–—]?\s*(.+?)\s*$', texto, re.IGNORECASE | re.MULTILINE )
            jefe = re.search( r'j\s*e\s*f\s*e\s*de\s*servicio:\s*([^\n]+)', texto, re.IGNORECASE)
            oficial = re.search(r'(?:Oficial\s*+de\s*+Serv[ií]cio:\s*|Oficial\s*+de\s*+Serv[ií]cio\s*)([^\n]+)', texto, re.IGNORECASE)
            operador = re.search(
                                    r'(?:'
                                    r'Operador\s*de\s*C[aá]mara'
                                    r'|Operador\s*C[aá]mara'
                                    r'|Op(?:erador)?\s*C[aá]mara'
                                    r'|Aux(?:iliar)?\s+Operador\s+de\s+C[aá]mara'
                                    r'|Aux\s*'
                                    r')'
                                    r'\s*:?\s*([^\n]+)',   # ← capturar solo la MISMA línea
                                    texto,
                                    re.IGNORECASE
                                )

            sae = re.search(r'(?:SAE\s*Nro.\s*|SAE:\s*°\s*|SAE\s*nº\s*|sae\s*|suceso\s*:\s*|suceso\s*|cad\s*:\s*|SAECAD\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*|sae\s*+nro.|carta\s*:|carta\s*n\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*|SAE\s*N.\s*º\s*|N°\s*|N°\s*:\s*)\s*(\d{8})', texto, re.IGNORECASE)
            resultado = re.search(r'\bresultado\b\s*[:\-–—]?\s*([^\n]+)', texto, re.IGNORECASE)
            resultado_texto = resultado.group(1).strip() if resultado else "S/D"            
            
            if operador:
                op_texto = operador.group(1).strip()

            # si empieza con "breve", "reseña", "resultado", etc → NO ES OPERADOR
                if re.match(r'^(breve|rese[nñ]a|resultado|fecha|hora)\b', op_texto, re.IGNORECASE):
                    operador_valor = "S/D"
                else:
                    operador_valor = op_texto
            else:
                operador_valor = "S/D"
            
            # Condición para detenidos
            detenido_valor = ""
            contraventor_valor = ""
            resultado_final = resultado_texto
            mapa_numeros = {
                            "un": 1, "uno": 1, "una": 1,
                            "dos": 2, "tres": 3, "cuatro": 4,
                            "cinco": 5, "seis": 6, "siete": 7,
                            "ocho": 8, "nueve": 9, "diez": 10
                            }

            def extraer_numero_resultado(texto):
                t = texto.lower()

                m = re.search(r'\b(\d+)\s*\(', t)
                if m:
                    return int(m.group(1))

                m = re.search(r'\b(\d+)\b', t)
                if m:
                    return int(m.group(1))

                m = re.search(r'\((\d+)\)', t)
                if m:
                    return int(m.group(1))

                for palabra, valor in mapa_numeros.items():
                    if re.search(rf'\b{palabra}\b', t):
                        return valor

                return ""

            # --- aplicar ---
            detenido_valor = contraventor_valor = ""

            if re.search(r'\bdeten', resultado_texto, re.IGNORECASE):
                detenido_valor = extraer_numero_resultado(resultado_texto)

            if re.search(r'\bcontra', resultado_texto, re.IGNORECASE):
                contraventor_valor = extraer_numero_resultado(resultado_texto)

            resultado_final = "" if detenido_valor or contraventor_valor else resultado_texto

            
            # Obtener y formatear la fecha
            fecha_texto = fecha.group(1).strip() if fecha else "S/D"
            # Antes: búsqueda rígida de fecha. Ahora se captura el texto después de 'Fecha:' y se parsea.
            fecha_match = re.search(r'Fecha\s*:\s*(.+?)(?:\s+Hora\b|\n|$)', texto, re.IGNORECASE)
            if fecha_match:
                fecha_texto = fecha_match.group(1).strip()
            else:
                # Fallback: buscar patrones de fecha en cualquier parte del texto
                fecha_busqueda = re.search(
                    r'(\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+(?:\s+de(?:l)?\s+\d{2,4})?|\d{1,2}[\/\-]\d{1,2}(?:[\/\-]\d{2,4})?|\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2})',
                    texto,
                    re.IGNORECASE
                )
                fecha_texto = fecha_busqueda.group(1).strip() if fecha_busqueda else "S/D"
            fecha_formateada = formatear_fecha(fecha_texto)
                
            datos.append({
                "RUTA": ruta,
                "FECHA": fecha_formateada,
                "TIPIFICACION": "",
                "CAUSA": causa.group(1).strip() if causa else "S/D",
                "SAE": int(sae.group(1).strip()) if sae else "S/D",
                "HORA": hora_final,
                "OPERADOR DE CAMARA": operador_valor,
                "OFICIAL DE SERVICIO": oficial.group(1).strip() if oficial else "S/D",
                "CONTRAVENTOR": contraventor_valor,
                "DETENIDO": detenido_valor,
                "JEFE DE SERVICIO": jefe.group(1).strip() if jefe else "S/D",
                "RESULTADO": resultado_final
            })
            # Construir bloque combinado para Word
            resena_bloque = re.search(
                r'(?:rese[nñ]a|rese[nñ]a)\s*:?\s*(.*?)Resultado\s*:',
                texto, re.IGNORECASE | re.DOTALL
            )
            if resena_bloque:
                resena_limpia = re.sub(r'\s+', ' ', resena_bloque.group(1)).strip()
            else:
                resena_limpia = "S/D"

            celda_combinada = (
                f"Fecha: {fecha_formateada}\n"
                f"Reseña: {resena_limpia}\n"
                f"Resultado: {resultado.group(1).strip() if resultado else 'S/D'}"
            )

            datos_2.append(celda_combinada)

        # --- Acumulación en Excel ---
        try:
            df_existente = pd.read_excel("contenido_archivos.xlsx")
            df_nuevo = pd.DataFrame(datos)
            df_final = pd.concat([df_existente, df_nuevo], ignore_index=True)
        except FileNotFoundError:
            df_final = pd.DataFrame(datos)
        
        # Guardar el DataFrame en Excel
        df_final.to_excel("contenido_archivos.xlsx", index=False)

        # Configurar formato de fecha en la columna B (columna 2) del Excel
        try:
            wb = openpyxl.load_workbook("contenido_archivos.xlsx")
            ws = wb.active
            if ws is None:
                raise ValueError("No active worksheet found in workbook")
            columnas = ["B", "E", "F", "I", "J"]
            columnas_limpiar = ['G', 'H', 'K']

            # Palabras a eliminar (sin signos)
            palabras_limpiar = [
                "personal", "contratado", "auxiliar",
                "pers", "contr", "aux", "cont",
                "oficial", "of",
                "operador", "oper",
                "lp", "primero", "mayor", "inspector",
                "subrio", "subcomisario"
            ]

            # Signos a eliminar siempre
            signos_eliminar = r"[.\-:;]"

            # Regex para palabras (tolerante a punto final)
            patrones = [
                re.compile(r'\b' + re.escape(p) + r'\.?\b', flags=re.IGNORECASE)
                for p in palabras_limpiar
            ]

            for col in columnas_limpiar:
                for celda in ws[col]:
                    if celda.row == 1:
                        continue

                    valor = celda.value
                    if not valor:
                        continue

                    texto = str(valor)

                    # 1) eliminar palabras específicas
                    for pat in patrones:
                        texto = pat.sub('', texto)

                    # 2) eliminar signos aislados (., -, :, ;)
                    texto = re.sub(signos_eliminar, ' ', texto)

                    # 3) eliminar cualquier número
                    texto = re.sub(r'\d+', ' ', texto)

                    # 4) limpiar espacios duplicados
                    texto = re.sub(r'\s+', ' ', texto).strip()

                    # 5) guardar
                    celda.value = texto if texto else "S/D"


                                
                for col in columnas:
                    for celda in ws[col]:
                        celda.alignment = Alignment(horizontal="center")

            # Establecer formato de fecha para la columna B
            for row in range(2, ws.max_row + 1):
                celda = ws.cell(row=row, column=2)  # Columna B = FECHA
                if celda.value and celda.value != "S/D":
                    try:
                        # Convertir a objeto fecha de Excel si es posible
                        if isinstance(celda.value, str):
                            fecha_obj = datetime.strptime(celda.value, "%d/%m/%Y")
                            celda.value = fecha_obj
                        celda.number_format = "DD/MM/AAAA"
                    except:
                        pass
            # Agregar hipervínculos en la columna "RUTA"
            col_titulo = 1  # La columna "RUTA" es la primera
            for row in range(2, ws.max_row + 1):
                ruta_pdf = ws.cell(row=row, column=col_titulo).value
                if ruta_pdf and isinstance(ruta_pdf, str) and ruta_pdf.lower().endswith(".pdf"):
                    ws.cell(row=row, column=col_titulo).hyperlink = Hyperlink(ref=f"A{row}", target=str(ruta_pdf))
                    ws.cell(row=row, column=col_titulo).font = Font(
                        color="0000FF", underline="single")

            wb.save("contenido_archivos.xlsx")
        except Exception as e:
            print(f"Error al formatear Excel: {e}")

        # --- Acumulación en Word ---
        try:
            doc = Document("reseña_archivos.docx")
        except Exception:
            doc = Document()
            
            # 2. Agrega un título (opcional)
            
            encabezado = doc.add_heading("División CENTROS DE MONITOREO URBANO", 0)
            encabezado.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            encabezado.runs[0].font.size = Pt(20)

# 3. Crea una tabla con un número de filas y columnas
            table = doc.add_table(rows=10, cols=5)
            table.style = 'Table Grid'
            
            row = table.rows[0]
            row.cells[0].merge(row.cells[4]) # Combina celda 0 con celda 4
            
            table.cell(0, 0).text = 'PARTE SEMANAL DEL XX/XX/2025 AL XX/XX/2025'
            table.cell(1, 0).text = 'MES DE DICIEMBRE/DIA'
            table.cell(1, 1).text = 'EVENTOS DEL PERÍODO'
            table.cell(1, 2).text = 'EVENTOS POSITIVOS'
            table.cell(1, 3).text = 'IMPUTADOS'
            table.cell(1, 4).text = 'EVENTOS RELEVANTES'
            # 5. Formato del título dentro de la celda combinada
            celda = table.cell(0, 0).paragraphs[0]
            celda.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = celda.runs[0]
            run.font.size = Pt(14)
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)


# 2. Acceder a una celda específica (ej. celda en fila 0, columna 0)
            #cell = table.cell(9, 0)
            fondo = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
            table.rows[0].cells[0]._tc.get_or_add_tcPr().append(fondo)


            # 6. Formato para el encabezado de la fila 1 (5 columnas → índice 0 a 4)
            for i in range(5):
                cell = table.cell(1, i)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = paragraph.runs[0]
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.bold = True

            res_encabezado = doc.add_heading('\n\nReseñas extraídas', 0)
            res_encabezado.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            res_encabezado.runs[0].font.size = Pt(14)
            
        for item in datos_2:
            doc.add_paragraph(item)
            doc.add_paragraph(
                '--------------------------------------------------------------')
        doc.save("reseña_archivos.docx")
    
        messagebox.showinfo("Éxito", "Datos extraídos y guardados en contenido_archivos.xlsx, reseña_archivos.docx")
        print("Datos extraídos y guardados en contenido_archivos.xlsx, reseña_archivos.docx")

        # Limpiar la interfaz después de ejecutar
        ventana_pegado.config(text="Haz clic en 'Seleccionar Archivos' para comenzar")
        
    except Exception as e:
        print(f"Error en extracción y conversión: {e}")
        traceback.print_exc()
        messagebox.showerror("Error", f"Error durante el procesamiento: {str(e)}")


# Crear ventana principal
root = tk.Tk()
root.title("Extracción de datos .PDF a: -Parte semanal-Detenidos Oficiales-")
root.config(bg="#822121")
root.geometry("840x400")

frame = tk.Frame(root)
frame.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)

ventana_pegado = tk.Label(
    frame,
    text="Haz clic en 'Seleccionar Archivos' para comenzar",
    bg="#1a1a1a",
    fg="#00FF00",
    font=("Helvetica", 12),
    relief="sunken",
    height=6
)
ventana_pegado.pack(fill=tk.BOTH, expand=True)

# Botones
boton_seleccionar = tk.Button(frame, text="📁 Seleccionar Archivos", width=20, height=2, command=seleccionar_archivos)
boton_seleccionar.config(background="#33AAFF", fg="white", font=("Helvetica", 10, "bold"))
boton_seleccionar.pack(side=tk.TOP, pady=5)

boton_ejecutar = tk.Button(frame, text="▶ Ejecutar", width=20, height=2, command=extraer_convertir)
boton_ejecutar.config(background="#33FF41", fg="white", font=("Helvetica", 10, "bold"))
boton_ejecutar.pack(side=tk.LEFT)

boton_cerrar = tk.Button(frame, text="❌ Cerrar", width=20, height=2, command=root.quit)
boton_cerrar.config(background="#FF3333", fg="white", font=("Helvetica", 10, "bold"))
boton_cerrar.pack(side=tk.RIGHT)

if __name__ == "__main__":
    root.mainloop()