import tkinter as tk
from tkinter import filedialog, messagebox
from pypdf import PdfReader
import re
import pandas as pd
from docx import Document
import openpyxl
from openpyxl.styles import Font, Alignment
import traceback
from datetime import datetime

# Variable global para almacenar rutas
rutas_global = []

def formatear_fecha(fecha_texto):
    """
    Convierte diferentes formatos de fecha al formato dd/mm/aaaa
    """
    if not fecha_texto or fecha_texto == "S/D":
        return "S/D"
    
    try:
        # Limpiar el texto de la fecha
        fecha_limpia = fecha_texto.strip()
        
        # Mapeo de meses en español a inglés
        meses_espanol = {
            'enero': 'January', 'febrero': 'February', 'marzo': 'March',
            'abril': 'April', 'mayo': 'May', 'junio': 'June',
            'julio': 'July', 'agosto': 'August', 'septiembre': 'September',
            'octubre': 'October', 'noviembre': 'November', 'diciembre': 'December'
        }
        
        # Patrones comunes de fecha
        patrones = [
            # Formato: "23 de octubre del 2025"
            r'(\d{1,2})\s+de\s+(\w+)\s+del?\s+(\d{4})',
            # Formato: "23 octubre 2025"
            r'(\d{1,2})\s+(\w+)\s+(\d{4})',
            # Formato: "23/10/2025" o "23-10-2025"
            r'(\d{1,2})[/-](\d{1,2})[/-](\d{4})',
            # Formato: "2025-10-23" (ISO)
            r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})'
        ]
        
        for patron in patrones:
            match = re.search(patron, fecha_limpia, re.IGNORECASE)
            if match:
                if patron == patrones[0] or patron == patrones[1]:
                    # Formato con mes en texto
                    dia, mes_texto, año = match.groups()
                    mes_texto = mes_texto.lower()
                    if mes_texto in meses_espanol:
                        fecha_obj = datetime.strptime(f"{dia} {meses_espanol[mes_texto]} {año}", "%d %B %Y")
                        return fecha_obj.strftime("%d/%m")
                
                elif patron == patrones[2]:
                    # Formato dd/mm/aaaa o dd-mm-aaaa
                    dia, mes, año = match.groups()
                    fecha_obj = datetime.strptime(f"{dia}/{mes}/{año}", "%d/%m/%Y")
                    return fecha_obj.strftime("%d/%m")
                
                elif patron == patrones[3]:
                    # Formato aaaa-mm-dd
                    año, mes, dia = match.groups()
                    fecha_obj = datetime.strptime(f"{año}-{mes}-{dia}", "%Y-%m-%d")
                    return fecha_obj.strftime("%d/%m")
        
        # Si no coincide con ningún patrón, intentar parseo directo
        try:
            fecha_obj = datetime.strptime(fecha_limpia, "%d/%m/%Y")
            return fecha_obj.strftime("%d/%m/%Y")
        except:
            try:
                fecha_obj = datetime.strptime(fecha_limpia, "%Y-%m-%d")
                return fecha_obj.strftime("%d/%m/%Y")
            except:
                # Si no se puede convertir, devolver el texto original
                return fecha_limpia
                
    except Exception as e:
        print(f"Error al formatear fecha '{fecha_texto}': {e}")
        return fecha_texto


def seleccionar_archivos():
    """Seleccionar archivos PDF usando diálogo de archivos"""
    try:
        global rutas_global
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")]
        )
        
        if archivos:
            rutas_global = list(archivos)
            archivos_texto = "Archivos seleccionados:\n" + "\n".join([f.split("/")[-1] for f in archivos])
            ventana_pegado.config(text=archivos_texto)
            print(f"Archivos seleccionados: {archivos}")
        else:
            ventana_pegado.config(text="No se seleccionaron archivos")
    except Exception as e:
        print(f"Error en selección de archivos: {e}")
        traceback.print_exc()
        messagebox.showerror("Error", f"Error al seleccionar archivos: {str(e)}")

def extraer_hora(texto):
    """
    Busca una hora en formato HH:MM solo si aparece
    después de 'Breve reseña:' o 'Breve reseña;'.
    Devuelve la hora encontrada o None.
    """

    # Expresión regular:
    # - Detecta "Breve reseña:" o "Breve reseña;"
    # - Captura la primera hora HH:MM después de esa frase
    patron = r"Breve reseña[:;].*?(\b[0-2][0-9]:[0-5][0-9]\b)"

    coincidencia = re.search(patron, texto, re.IGNORECASE | re.DOTALL)

    if coincidencia:
        return coincidencia.group(1)  # La hora
    return None

def extraer_convertir():
    """Extraer datos de archivos PDF y guardarlos en Excel y Word"""
    try:
        if not rutas_global:
            messagebox.showwarning("Advertencia", "Por favor, selecciona archivos PDF primero.")
            return
            
        datos, datos_2 = [], []
        contenido = ""
        error = ""
        
        

        for ruta in rutas_global:
            try:
                documento = PdfReader(ruta)
                contenido += "----------" + ruta + "----------------\n"
                for pagina in documento.pages:
                    texto_pagina = pagina.extract_text()
                    if texto_pagina:
                        contenido += texto_pagina + "\n"
            except Exception as e:
                contenido += f"Error al leer {ruta}: {str(e)}\n"

        bloques = re.split(r'-{10,}(.+?)-{10,}\n', contenido)
        for i in range(1, len(bloques), 2):
            
            ruta = bloques[i].strip()
            texto = bloques[i+1].strip() if i+1 < len(bloques) else ""
            fecha = re.search(r'Fecha\s*:\s*('r'\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}'r'|\d{1,2}[/-]\d{1,2}[/-]\d{4}'r'|\d{4}[/-]\d{1,2}[/-]\d{1,2}'r')', texto,re.IGNORECASE)
            patron_breve = r"Breve reseña[:;].*?(\b[0-2][0-9]:[0-5][0-9]\b)"############################
            coincidencia_breve = re.search(patron_breve, texto, re.IGNORECASE | re.DOTALL)#########################
            if coincidencia_breve:
                    hora = coincidencia_breve.group(1)
            else:
            # Si no está en reseña breve, buscar una hora estándar
                hora_busqueda = re.search(r'(Hora|Horario)\s*[:\-]?\s*(\d{1,2}:\d{2})(?:\s*hs?)?', texto, re.IGNORECASE)
                    #re.search(r'(Hora|Horario)\s*[:\-]?\s*(\d{1,2}:\d{2})(?:\s*hs?)?', texto, re.IGNORECASE)
            causa = re.search(r'^\s*(?:causa|hecho)\b\s*[\s*:\-–—]?\s*(.+?)\s*$', texto, re.IGNORECASE | re.MULTILINE )
            jefe = re.search(r'Jefe\s*de\s*Servicio:\s*([^\n]+)', texto, re.IGNORECASE)
            oficial = re.search(r'(?:Oficial\s*+de\s*+Serv[ií]cio:\s*|Oficial\s*+de\s*+Serv[ií]cio\s*)([^\n]+)', texto, re.IGNORECASE)
            operador = re.search(r'Operador\s*+de\s*+C[aá]mara\s*(?:Aux)?\s*:?\s*([^\n]+)', texto, re.IGNORECASE)          
            sae = re.search(r'(?:SAE\s*Nro.\s*|SAE\s*nº\s*|sae\s*|suceso\s*:\s*|suceso\s*|cad\s*:\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*|sae\s*+nro.|carta\s*:|carta\s*n\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*|SAE\s*N.\s*º\s*|N°\s*|N°\s*:\s*)\s*(\d{8})', texto, re.IGNORECASE)
            resultado = re.search(r'\bresultado\b\s*[:\-–—]?\s*([^\n]+)', texto, re.IGNORECASE)
            resultado_texto = resultado.group(1).strip() if resultado else "S/D"
            
            # Condición para detenidos
            try:
                if re.search(r'\s*det', resultado_texto, re.IGNORECASE):
                    num = re.search(r'\((\d+)\)', resultado_texto)
                    detenido_valor = str(int(num.group(1)))
                else:
                    detenido_valor = ""
                
                if re.search(r'\s*cont', resultado_texto, re.IGNORECASE):
                    num = re.search(r'\((\d+)\)', resultado_texto)
                    contraventor_valor = int(num.group(1))
                else:
                    contraventor_valor = ""
                
                if detenido_valor or contraventor_valor:
                    resultado_final = ""
                else:
                    resultado_final = resultado_texto
            except Exception as e:
                error += f"Error al leer {ruta} {str(e)}\n"
                print(error)
            
            
            # Obtener y formatear la fecha
            fecha_texto = fecha.group(1).strip() if fecha else "S/D"
            fecha_formateada = formatear_fecha(fecha_texto)
                
            datos.append({
                "RUTA": ruta,
                "FECHA": fecha_formateada,  # Usar fecha formateada
                "TIPIFICACION": "",
                "CAUSA": causa.group(1).strip() if causa else "S/D",
                "SAE": int(sae.group(1).strip()) if sae else "S/D",
                "HORA": hora_busqueda.group(2).strip() if hora_busqueda else "S/D",
                "OPERADOR DE CAMARA": operador.group(1).strip() if operador else "S/D",
                "OFICIAL DE SERVICIO": oficial.group(1).strip() if oficial else "S/D",
                "CONTRAVENTOR": contraventor_valor,
                "DETENIDO": detenido_valor,
                "JEFE DE SERVICIO": jefe.group(1).strip() if jefe else "S/D",
                "RESULTADO": resultado_final
            })

            resena_bloque = re.search(
                r'(?:rese[nñ]a|rese[nñ]a)\s*:?\s*(.*?)Resultado\s*:',
                texto, re.IGNORECASE | re.DOTALL
            )
            if resena_bloque:
                resena_limpia = re.sub(r'\s+', ' ', resena_bloque.group(1)).strip()
            else:
                resena_limpia = "S/D"

            celda_combinada = (
                f"Fecha: {fecha_formateada}\n"  # Usar fecha formateada aquí también
                f"Reseña: {resena_limpia}\n"
                f"Resultado: {resultado.group(1).strip() if resultado else 'S/D'}"
            )

            datos_2.append(celda_combinada)

        # --- Acumulación en Excel ---
        try:
            df_existente = pd.read_excel("contenido_archivos.xlsx")
            df_nuevo = pd.DataFrame(datos)
            df_final = pd.concat([df_existente, df_nuevo], ignore_index=True)
        except FileNotFoundError:
            df_final = pd.DataFrame(datos)
        
        # Guardar el DataFrame en Excel
        df_final.to_excel("contenido_archivos.xlsx", index=False)

        # Configurar formato de fecha en la columna B (columna 2) del Excel
        try:
            wb = openpyxl.load_workbook("contenido_archivos.xlsx")
            ws = wb.active
            columnas = ["B", "E", "F", "I", "J"]
            
            for col in columnas:
                for celda in ws[col]:
                    celda.alignment = Alignment(horizontal="center")

            # Establecer formato de fecha para la columna B
            for row in range(2, ws.max_row + 1):
                celda = ws.cell(row=row, column=2)  # Columna B = FECHA
                if celda.value and celda.value != "S/D":
                    try:
                        # Convertir a objeto fecha de Excel si es posible
                        if isinstance(celda.value, str):
                            fecha_obj = datetime.strptime(celda.value, "%d/%m/%Y")
                            celda.value = fecha_obj
                        celda.number_format = "DD/MM/AAAA"
                    except:
                        pass
            
            # Agregar hipervínculos en la columna "RUTA"
            col_titulo = 1  # La columna "RUTA" es la primera
            for row in range(2, ws.max_row + 1):
                ruta_pdf = ws.cell(row=row, column=col_titulo).value
                if ruta_pdf and ruta_pdf.lower().endswith(".pdf"):
                    ws.cell(row=row, column=col_titulo).hyperlink = ruta_pdf
                    ws.cell(row=row, column=col_titulo).font = Font(
                        color="0000FF", underline="single")

            wb.save("contenido_archivos.xlsx")
        except Exception as e:
            print(f"Error al formatear Excel: {e}")

        # --- Acumulación en Word ---
        try:
            doc = Document("reseña_archivos.docx")
        except Exception:
            doc = Document()
            doc.add_heading('Reseñas extraídas', 0)
        for item in datos_2:
            doc.add_paragraph(item)
            doc.add_paragraph(
                '--------------------------------------------------------------')
        doc.save("reseña_archivos.docx")
    
        messagebox.showinfo("Éxito", "Datos extraídos y guardados en contenido_archivos.xlsx, reseña_archivos.docx")
        print("Datos extraídos y guardados en contenido_archivos.xlsx, reseña_archivos.docx")

        # Limpiar la interfaz después de ejecutar
        ventana_pegado.config(text="Haz clic en 'Seleccionar Archivos' para comenzar")
        
    except Exception as e:
        print(f"Error en extracción y conversión: {e}")
        traceback.print_exc()
        messagebox.showerror("Error", f"Error durante el procesamiento: {str(e)}")


# Crear ventana principal
root = tk.Tk()
root.title("Extracción de datos .PDF a: -Parte semanal-Detenidos Oficiales-")
root.config(bg="#822121")
root.geometry("840x400")

frame = tk.Frame(root)
frame.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)

ventana_pegado = tk.Label(
    frame,
    text="Haz clic en 'Seleccionar Archivos' para comenzar",
    bg="#2b74be",
    font=("Helvetica", 14),
    relief="sunken",
    fg="white",
    height=6
)
ventana_pegado.pack(fill=tk.BOTH, expand=True)

# Botones
boton_seleccionar = tk.Button(frame, text="Seleccionar Archivos", width=20, height=2, command=seleccionar_archivos)
boton_seleccionar.config(background="#33AAFF", fg="white", activeforeground="blue")
boton_seleccionar.pack(side=tk.TOP, pady=5)

boton_ejecutar = tk.Button(frame, text="Ejecutar", width=20, height=2, command=extraer_convertir)
boton_ejecutar.config(background="#33FF41", fg="darkblue", activeforeground="green")
boton_ejecutar.pack(side=tk.LEFT)

boton_cerrar = tk.Button(frame, text="Cerrar", width=20, height=2, command=root.quit)
boton_cerrar.config(background="#FF3333", fg="darkblue", activeforeground="red")
boton_cerrar.pack(side=tk.RIGHT)

if __name__ == "__main__":
    root.mainloop()