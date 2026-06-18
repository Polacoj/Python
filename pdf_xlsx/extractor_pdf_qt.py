"""
Extractor de datos desde archivos PDF → Excel y Word
Interfaz construida con PyQt5
"""

import re
import sys
import traceback
from datetime import datetime

import pandas as pd
import openpyxl
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Alignment
from openpyxl.worksheet.hyperlink import Hyperlink

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QTextEdit, QFileDialog, QMessageBox,
    QProgressBar, QFrame, QSizePolicy,
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor, QTextCursor


# ──────────────────────────────────────────────────────────────
#  Helper de fechas (sin cambios)
# ──────────────────────────────────────────────────────────────

def formatear_fecha(fecha_texto):
    if not fecha_texto or fecha_texto == "S/D":
        return "S/D"
    try:
        fecha_limpia = fecha_texto.strip()
        meses = {
            "enero": 1, "ene": 1, "febrero": 2, "feb": 2,
            "marzo": 3, "mar": 3, "abril": 4, "abr": 4,
            "mayo": 5, "may": 5, "junio": 6, "jun": 6,
            "julio": 7, "jul": 7, "agosto": 8, "ago": 8,
            "septiembre": 9, "sep": 9, "setiembre": 9,
            "octubre": 10, "oct": 10, "noviembre": 11, "nov": 11,
            "diciembre": 12, "dic": 12,
        }
        m = re.search(r"(\d{1,2})\s+de\s+([A-záéíóúñ]+)", fecha_limpia, re.IGNORECASE)
        if not m:
            m = re.search(r"(\d{1,2})\s+([A-záéíóúñ]+)", fecha_limpia, re.IGNORECASE)
        if m:
            mes = meses.get(m.group(2).lower().replace(".", ""))
            if mes:
                return f"{int(m.group(1)):02d}/{mes:02d}"
        m = re.search(r"(\d{1,2})[\/\-](\d{1,2})(?:[\/\-]\d{2,4})?", fecha_limpia)
        if m:
            return f"{int(m.group(1)):02d}/{int(m.group(2)):02d}"
        m = re.search(r"(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})", fecha_limpia)
        if m:
            return f"{int(m.group(3)):02d}/{int(m.group(2)):02d}"
        for fmt in ("%d/%m/%Y", "%d/%m", "%Y-%m-%d", "%m-%d"):
            try:
                return datetime.strptime(fecha_limpia, fmt).strftime("%d/%m")
            except Exception:
                continue
        return fecha_limpia
    except Exception as e:
        print(f"Error al formatear fecha '{fecha_texto}': {e}")
        return fecha_texto


# ──────────────────────────────────────────────────────────────
#  Worker: corre en hilo separado para no congelar la UI
# ──────────────────────────────────────────────────────────────

class Worker(QThread):
    log      = pyqtSignal(str, str)   # (mensaje, color_hex)
    progress = pyqtSignal(int)        # 0-100
    status   = pyqtSignal(str)
    finished = pyqtSignal(bool, str)  # (éxito, mensaje)

    def __init__(self, rutas: list[str]):
        super().__init__()
        self.rutas = rutas

    def run(self):
        try:
            datos, datos_2 = [], []

            for idx, ruta in enumerate(self.rutas, 1):
                nombre = ruta.split("/")[-1] if "/" in ruta else ruta.split("\\")[-1]
                self.log.emit(f"\n🔍  [{idx}/{len(self.rutas)}]  {nombre}", "#e2e8f0")
                self.status.emit(f"Procesando {idx}/{len(self.rutas)}: {nombre}")
                self.progress.emit(int((idx - 1) / len(self.rutas) * 100))

                try:
                    doc_pdf = PdfReader(ruta)
                    texto = "\n".join(p.extract_text() or "" for p in doc_pdf.pages)
                except Exception as ex:
                    self.log.emit(f"❌  No se pudo leer: {ex}", "#f87171")
                    texto = ""

                # ── Regex (idénticas al original) ──────────────────────
                hora_resena = re.search(
                    r"(breve\s+reseña|reseña)\s*[:\-]?\s*(.*?)\b(?:siendo\s+las|a\s+las)\s*(\d{1,2}:\d{2})",
                    texto, re.IGNORECASE | re.DOTALL,
                )
                hora_general = re.search(
                    r"(hora|horario)\s*[:\-]?\s*(\d{1,2}:\d{2})", texto, re.IGNORECASE
                )
                hora_final = (hora_resena.group(3) if hora_resena
                              else hora_general.group(2) if hora_general
                              else "S/D")

                causa    = re.search(r"^\s*(?:causa|hecho)\b\s*[\s*:\-–—]?\s*(.+?)\s*$", texto, re.IGNORECASE | re.MULTILINE)
                jefe     = re.search(r"j\s*e\s*f\s*e\s*de\s*servicio:\s*([^\n]+)", texto, re.IGNORECASE)
                oficial  = re.search(r"(?:Oficial\s*+de\s*+Serv[ií]cio:\s*|Oficial\s*+de\s*+Serv[ií]cio\s*)([^\n]+)", texto, re.IGNORECASE)
                operador = re.search(
                    r"(?:Operador\s*de\s*C[aá]mara|Operador\s*C[aá]mara|Op(?:erador)?\s*C[aá]mara"
                    r"|Aux(?:iliar)?\s+Operador\s+de\s+C[aá]mara|Aux\s*)\s*:?\s*([^\n]+)",
                    texto, re.IGNORECASE,
                )
                sae = re.search(
                    r"(?:SAE\s*Nro.\s*|SAE:\s*°\s*|SAE\s*nº\s*|sae\s*|suceso\s*:\s*|suceso\s*"
                    r"|cad\s*:\s*|SAECAD\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*"
                    r"|sae\s*+nro.|carta\s*:|carta\s*n\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*"
                    r"|SAE\s*N.\s*º\s*|N°\s*|N°\s*:\s*)\s*(\d{8})",
                    texto, re.IGNORECASE,
                )
                resultado       = re.search(r"\bresultado\b\s*[:\-–—]?\s*([^\n]+)", texto, re.IGNORECASE)
                resultado_texto = resultado.group(1).strip() if resultado else "S/D"

                if operador:
                    op_texto = operador.group(1).strip()
                    operador_valor = (
                        "S/D"
                        if re.match(r"^(breve|rese[nñ]a|resultado|fecha|hora)\b", op_texto, re.IGNORECASE)
                        else op_texto
                    )
                else:
                    operador_valor = "S/D"

                mapa_numeros = {
                    "un": 1, "uno": 1, "una": 1, "dos": 2, "tres": 3,
                    "cuatro": 4, "cinco": 5, "seis": 6, "siete": 7,
                    "ocho": 8, "nueve": 9, "diez": 10,
                }

                def extraer_numero(t):
                    t = t.lower()
                    for pat in (r"\b(\d+)\s*\(", r"\b(\d+)\b", r"\((\d+)\)"):
                        m = re.search(pat, t)
                        if m:
                            return int(m.group(1))
                    for palabra, valor in mapa_numeros.items():
                        if re.search(rf"\b{palabra}\b", t):
                            return valor
                    return ""

                detenido_valor = contraventor_valor = ""
                if re.search(r"\bdeten", resultado_texto, re.IGNORECASE):
                    detenido_valor = extraer_numero(resultado_texto)
                if re.search(r"\bcontra", resultado_texto, re.IGNORECASE):
                    contraventor_valor = extraer_numero(resultado_texto)
                resultado_final = "" if (detenido_valor or contraventor_valor) else resultado_texto

                fecha_match = re.search(r"Fecha\s*:\s*(.+?)(?:\s+Hora\b|\n|$)", texto, re.IGNORECASE)
                if fecha_match:
                    fecha_texto_raw = fecha_match.group(1).strip()
                else:
                    fb = re.search(
                        r"(\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+(?:\s+de(?:l)?\s+\d{2,4})?|\d{1,2}[\/\-]\d{1,2}(?:[\/\-]\d{2,4})?|\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2})",
                        texto, re.IGNORECASE,
                    )
                    fecha_texto_raw = fb.group(1).strip() if fb else "S/D"
                fecha_fmt = formatear_fecha(fecha_texto_raw)

                datos.append({
                    "RUTA": ruta,
                    "FECHA": fecha_fmt,
                    "TIPIFICACION": "",
                    "CAUSA": causa.group(1).strip() if causa else "S/D",
                    "SAE": int(sae.group(1).strip()) if sae else "S/D",
                    "HORA": hora_final,
                    "OPERADOR DE CAMARA": operador_valor,
                    "OFICIAL DE SERVICIO": oficial.group(1).strip() if oficial else "S/D",
                    "CONTRAVENTOR": contraventor_valor,
                    "DETENIDO": detenido_valor,
                    "JEFE DE SERVICIO": jefe.group(1).strip() if jefe else "S/D",
                    "RESULTADO": resultado_final,
                })

                resena_bloque = re.search(
                    r"(?:rese[nñ]a)\s*:?\s*(.*?)Resultado\s*:", texto, re.IGNORECASE | re.DOTALL
                )
                resena_limpia = (
                    re.sub(r"\s+", " ", resena_bloque.group(1)).strip()
                    if resena_bloque else "S/D"
                )
                datos_2.append(
                    f"Fecha: {fecha_fmt}\n"
                    f"Reseña: {resena_limpia}\n"
                    f"Resultado: {resultado.group(1).strip() if resultado else 'S/D'}"
                )
                self.log.emit(f"✅  Extraído — Fecha: {fecha_fmt}  SAE: {sae.group(1) if sae else 'S/D'}", "#4ade80")

            # ── Excel ───────────────────────────────────────────────
            self.log.emit("\n💾  Guardando Excel…", "#e2e8f0")
            try:
                df_existente = pd.read_excel("contenido_archivos.xlsx")
                df_final = pd.concat([df_existente, pd.DataFrame(datos)], ignore_index=True)
            except FileNotFoundError:
                df_final = pd.DataFrame(datos)
            df_final.to_excel("contenido_archivos.xlsx", index=False)

            try:
                wb = openpyxl.load_workbook("contenido_archivos.xlsx")
                ws = wb.active
                palabras_limpiar = [
                    "personal", "contratado", "auxiliar", "pers", "contr", "aux",
                    "cont", "oficial", "of", "operador", "oper", "lp", "primero",
                    "mayor", "inspector", "subrio", "subcomisario",
                ]
                patrones = [re.compile(r"\b" + re.escape(p) + r"\.?\b", re.IGNORECASE) for p in palabras_limpiar]

                for col in ["G", "H", "K"]:
                    for celda in ws[col]:
                        if celda.row == 1 or not celda.value:
                            continue
                        t = str(celda.value)
                        for pat in patrones:
                            t = pat.sub("", t)
                        t = re.sub(r"[.\-:;]", " ", t)
                        t = re.sub(r"\d+", " ", t)
                        t = re.sub(r"\s+", " ", t).strip()
                        celda.value = t or "S/D"

                for col in ["B", "E", "F", "I", "J"]:
                    for celda in ws[col]:
                        celda.alignment = Alignment(horizontal="center")

                for row in range(2, ws.max_row + 1):
                    celda = ws.cell(row=row, column=2)
                    if celda.value and celda.value != "S/D" and isinstance(celda.value, str):
                        try:
                            celda.value = datetime.strptime(celda.value, "%d/%m/%Y")
                            celda.number_format = "DD/MM/AAAA"
                        except Exception:
                            pass

                for row in range(2, ws.max_row + 1):
                    ruta_pdf = ws.cell(row=row, column=1).value
                    if ruta_pdf and isinstance(ruta_pdf, str) and ruta_pdf.lower().endswith(".pdf"):
                        ws.cell(row=row, column=1).hyperlink = Hyperlink(ref=f"A{row}", target=str(ruta_pdf))
                        ws.cell(row=row, column=1).font = Font(color="0000FF", underline="single")

                wb.save("contenido_archivos.xlsx")
            except Exception as ex:
                self.log.emit(f"⚠️  Error al formatear Excel: {ex}", "#fbbf24")

            # ── Word ────────────────────────────────────────────────
            self.log.emit("💾  Guardando Word…", "#e2e8f0")
            try:
                doc = Document("reseña_archivos.docx")
            except Exception:
                doc = Document()
                enc = doc.add_heading("División CENTROS DE MONITOREO URBANO", 0)
                enc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                enc.runs[0].font.size = Pt(20)

                table = doc.add_table(rows=10, cols=5)
                table.style = "Table Grid"
                fila0 = table.rows[0]
                fila0.cells[0].merge(fila0.cells[4])
                table.cell(0, 0).text = "PARTE SEMANAL DEL XX/XX/2025 AL XX/XX/2025"
                for i, txt in enumerate(["MES DE DICIEMBRE/DIA", "EVENTOS DEL PERÍODO",
                                          "EVENTOS POSITIVOS", "IMPUTADOS", "EVENTOS RELEVANTES"]):
                    table.cell(1, i).text = txt

                celda_p = table.cell(0, 0).paragraphs[0]
                celda_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = celda_p.runs[0]
                run.font.size = Pt(14)
                run.font.name = "Arial"
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                fondo = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls("w")))
                table.rows[0].cells[0]._tc.get_or_add_tcPr().append(fondo)

                for i in range(5):
                    p = table.cell(1, i).paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.runs[0]
                    r.font.size = Pt(10)
                    r.font.name = "Arial"
                    r.font.bold = True

                res_enc = doc.add_heading("\n\nReseñas extraídas", 0)
                res_enc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                res_enc.runs[0].font.size = Pt(14)

            for item in datos_2:
                doc.add_paragraph(item)
                doc.add_paragraph("--------------------------------------------------------------")
            doc.save("reseña_archivos.docx")

            self.progress.emit(100)
            self.log.emit("\n✅  Proceso completado.", "#4ade80")
            self.finished.emit(True, "Datos guardados en contenido_archivos.xlsx y reseña_archivos.docx")

        except Exception as ex:
            traceback.print_exc()
            self.log.emit(f"\n❌  {ex}", "#f87171")
            self.finished.emit(False, str(ex))


# ──────────────────────────────────────────────────────────────
#  Ventana principal
# ──────────────────────────────────────────────────────────────

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.rutas: list[str] = []
        self.worker: Worker | None = None
        self._build_ui()

    def _build_ui(self):
        self.setWindowTitle("Extractor PDF → Excel / Word")
        self.resize(860, 540)
        self.setMinimumSize(680, 420)

        # Fondo general oscuro
        self.setStyleSheet("""
            QMainWindow, QWidget#central {
                background-color: #1a1a2e;
            }
            QTextEdit {
                background-color: #0f172a;
                color: #cbd5e1;
                font-family: Menlo, Consolas, monospace;
                font-size: 12px;
                border: 1px solid #1e293b;
                border-radius: 6px;
                padding: 6px;
            }
            QLabel#title {
                color: white;
                font-size: 16px;
                font-weight: bold;
            }
            QLabel#subtitle {
                color: #64748b;
                font-size: 10px;
            }
            QLabel#status {
                color: #64748b;
                font-size: 11px;
                font-style: italic;
            }
            QProgressBar {
                background-color: #334155;
                border: none;
                border-radius: 3px;
                height: 6px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #4ade80;
                border-radius: 3px;
            }
            QPushButton {
                border-radius: 5px;
                padding: 8px 18px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton#btn_sel {
                background-color: #0369a1;
                color: white;
            }
            QPushButton#btn_sel:hover   { background-color: #0284c7; }
            QPushButton#btn_sel:pressed { background-color: #075985; }
            QPushButton#btn_sel:disabled { background-color: #374151; color: #6b7280; }

            QPushButton#btn_run {
                background-color: #16a34a;
                color: white;
            }
            QPushButton#btn_run:hover   { background-color: #15803d; }
            QPushButton#btn_run:pressed { background-color: #166534; }
            QPushButton#btn_run:disabled { background-color: #374151; color: #6b7280; }

            QPushButton#btn_clr {
                background-color: transparent;
                color: #94a3b8;
                border: 1px solid #334155;
            }
            QPushButton#btn_clr:hover { background-color: #1e293b; }

            QPushButton#btn_exit {
                background-color: transparent;
                color: #f87171;
                border: 1px solid #7f1d1d;
            }
            QPushButton#btn_exit:hover { background-color: #450a0a; }

            QFrame#separator {
                background-color: #1e293b;
            }
        """)

        central = QWidget()
        central.setObjectName("central")
        self.setCentralWidget(central)
        root_layout = QVBoxLayout(central)
        root_layout.setContentsMargins(0, 0, 0, 0)
        root_layout.setSpacing(0)

        # ── Header ────────────────────────────────────────────
        header = QWidget()
        header.setStyleSheet("background-color: #0f172a;")
        h_lay = QHBoxLayout(header)
        h_lay.setContentsMargins(16, 12, 16, 12)

        lbl_title = QLabel("📋  Extractor PDF → Excel / Word")
        lbl_title.setObjectName("title")
        lbl_subtitle = QLabel("División Centros de Monitoreo Urbano")
        lbl_subtitle.setObjectName("subtitle")

        title_col = QVBoxLayout()
        title_col.setSpacing(2)
        title_col.addWidget(lbl_title)
        title_col.addWidget(lbl_subtitle)
        h_lay.addLayout(title_col)
        h_lay.addStretch()
        root_layout.addWidget(header)
        root_layout.addWidget(self._sep())

        # ── Área de log ────────────────────────────────────────
        self.txt_log = QTextEdit()
        self.txt_log.setReadOnly(True)
        root_layout.addWidget(self.txt_log, stretch=1)

        # ── Barra de progreso ──────────────────────────────────
        self.progress = QProgressBar()
        self.progress.setValue(0)
        self.progress.setTextVisible(False)
        self.progress.setFixedHeight(6)
        root_layout.addWidget(self.progress)

        # ── Botones ────────────────────────────────────────────
        btn_row = QWidget()
        btn_row.setStyleSheet("background-color: #1a1a2e;")
        b_lay = QHBoxLayout(btn_row)
        b_lay.setContentsMargins(12, 8, 12, 8)
        b_lay.setSpacing(8)

        self.btn_sel = QPushButton("📁  Seleccionar PDFs")
        self.btn_sel.setObjectName("btn_sel")
        self.btn_sel.setCursor(Qt.PointingHandCursor)

        self.btn_run = QPushButton("▶  Ejecutar extracción")
        self.btn_run.setObjectName("btn_run")
        self.btn_run.setEnabled(False)
        self.btn_run.setCursor(Qt.PointingHandCursor)

        self.btn_clr = QPushButton("🗑  Limpiar log")
        self.btn_clr.setObjectName("btn_clr")
        self.btn_clr.setCursor(Qt.PointingHandCursor)

        self.btn_exit = QPushButton("✕  Cerrar")
        self.btn_exit.setObjectName("btn_exit")
        self.btn_exit.setCursor(Qt.PointingHandCursor)

        b_lay.addWidget(self.btn_sel)
        b_lay.addWidget(self.btn_run)
        b_lay.addWidget(self.btn_clr)
        b_lay.addStretch()
        b_lay.addWidget(self.btn_exit)
        root_layout.addWidget(btn_row)

        # ── Barra de estado ────────────────────────────────────
        root_layout.addWidget(self._sep())
        status_bar = QWidget()
        status_bar.setStyleSheet("background-color: #0f172a;")
        s_lay = QHBoxLayout(status_bar)
        s_lay.setContentsMargins(14, 4, 14, 4)
        self.lbl_status = QLabel("Listo.")
        self.lbl_status.setObjectName("status")
        s_lay.addWidget(self.lbl_status)
        root_layout.addWidget(status_bar)

        # ── Conexiones ─────────────────────────────────────────
        self.btn_sel.clicked.connect(self.seleccionar_archivos)
        self.btn_run.clicked.connect(self.iniciar_extraccion)
        self.btn_clr.clicked.connect(self.limpiar_log)
        self.btn_exit.clicked.connect(self.close)

        self.log("Hacé clic en '📁 Seleccionar PDFs' para comenzar.", "#64748b")

    # ── Helpers UI ─────────────────────────────────────────────

    def _sep(self):
        sep = QFrame()
        sep.setObjectName("separator")
        sep.setFixedHeight(1)
        return sep

    def log(self, msg: str, color: str = "#cccccc"):
        self.txt_log.setTextColor(QColor(color))
        self.txt_log.append(msg)
        self.txt_log.moveCursor(QTextCursor.End)

    def set_status(self, msg: str):
        self.lbl_status.setText(msg)

    def set_busy(self, busy: bool):
        self.btn_sel.setEnabled(not busy)
        self.btn_run.setEnabled(not busy)
        if busy:
            self.progress.setRange(0, 0)   # animación indeterminada
        else:
            self.progress.setRange(0, 100)

    # ── Acciones ───────────────────────────────────────────────

    def seleccionar_archivos(self):
        archivos, _ = QFileDialog.getOpenFileNames(
            self,
            "Seleccionar archivos PDF",
            "",
            "Archivos PDF (*.pdf);;Todos los archivos (*.*)",
        )
        if archivos:
            self.rutas = archivos
            self.txt_log.clear()
            self.log(f"✅  {len(self.rutas)} archivo(s) seleccionado(s):", "#4ade80")
            for r in self.rutas:
                nombre = r.split("/")[-1] if "/" in r else r.split("\\")[-1]
                self.log(f"   📄  {nombre}", "#94a3b8")
            self.btn_run.setEnabled(True)
            self.set_status(f"{len(self.rutas)} PDF(s) listo(s).")
        else:
            self.set_status("No se seleccionaron archivos.")

    def iniciar_extraccion(self):
        if not self.rutas:
            QMessageBox.warning(self, "Advertencia", "Seleccioná archivos PDF primero.")
            return
        self.set_busy(True)
        self.set_status("Procesando PDFs…")

        self.worker = Worker(self.rutas)
        self.worker.log.connect(self.log)
        self.worker.progress.connect(self._on_progress)
        self.worker.status.connect(self.set_status)
        self.worker.finished.connect(self._on_finished)
        self.worker.start()

    def _on_progress(self, value: int):
        self.progress.setRange(0, 100)
        self.progress.setValue(value)

    def _on_finished(self, exito: bool, mensaje: str):
        self.set_busy(False)
        self.progress.setRange(0, 100)
        self.progress.setValue(100 if exito else 0)
        self.set_status(mensaje)
        if exito:
            QMessageBox.information(self, "Éxito", mensaje)
        else:
            QMessageBox.critical(self, "Error", f"Error durante el procesamiento:\n{mensaje}")

    def limpiar_log(self):
        self.txt_log.clear()
        self.progress.setValue(0)
        self.log("Log limpiado.", "#64748b")
        self.set_status("Listo.")


# ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
