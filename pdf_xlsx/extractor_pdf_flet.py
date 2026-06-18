"""
Extractor de datos desde archivos PDF → Excel y Word
Interfaz construida con Flet
"""

import re
import threading
import traceback
from datetime import datetime

import pandas as pd
import openpyxl
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Alignment
from openpyxl.worksheet.hyperlink import Hyperlink

import flet as ft

# ──────────────────────────────────────────────────────────────
#  Helpers de lógica (sin cambios respecto al original)
# ──────────────────────────────────────────────────────────────


def formatear_fecha(fecha_texto):
    """Convierte distintos formatos de fecha y devuelve solo 'dd/mm'."""
    if not fecha_texto or fecha_texto == "S/D":
        return "S/D"
    try:
        fecha_limpia = fecha_texto.strip()
        meses_espanol = {
            "enero": 1,
            "ene": 1,
            "febrero": 2,
            "feb": 2,
            "marzo": 3,
            "mar": 3,
            "abril": 4,
            "abr": 4,
            "mayo": 5,
            "may": 5,
            "junio": 6,
            "jun": 6,
            "julio": 7,
            "jul": 7,
            "agosto": 8,
            "ago": 8,
            "septiembre": 9,
            "sep": 9,
            "setiembre": 9,
            "octubre": 10,
            "oct": 10,
            "noviembre": 11,
            "nov": 11,
            "diciembre": 12,
            "dic": 12,
        }

        m = re.search(r"(\d{1,2})\s+de\s+([A-záéíóúñ]+)", fecha_limpia, re.IGNORECASE)
        if not m:
            m = re.search(r"(\d{1,2})\s+([A-záéíóúñ]+)", fecha_limpia, re.IGNORECASE)
        if m:
            dia = int(m.group(1))
            mes = meses_espanol.get(m.group(2).lower().replace(".", ""))
            if mes:
                return f"{dia:02d}/{mes:02d}"

        m = re.search(r"(\d{1,2})[\/\-](\d{1,2})(?:[\/\-]\d{2,4})?", fecha_limpia)
        if m:
            return f"{int(m.group(1)):02d}/{int(m.group(2)):02d}"

        m = re.search(r"(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})", fecha_limpia)
        if m:
            return f"{int(m.group(3)):02d}/{int(m.group(2)):02d}"

        for fmt in ("%d/%m/%Y", "%d/%m", "%Y-%m-%d", "%m-%d"):
            try:
                return datetime.strptime(fecha_limpia, fmt).strftime("%d/%m")
            except Exception:
                continue

        return fecha_limpia
    except Exception as e:
        print(f"Error al formatear fecha '{fecha_texto}': {e}")
        return fecha_texto


# ──────────────────────────────────────────────────────────────
#  App Flet
# ──────────────────────────────────────────────────────────────


def main(page: ft.Page):
    page.title = "Extractor PDF → Excel / Word"
    page.window.width = 860
    page.window.height = 560
    page.window.min_width = 680
    page.theme_mode = ft.ThemeMode.DARK
    page.bgcolor = "#1a1a2e"
    page.padding = 0

    # ── Estado interno ─────────────────────────────────────────
    rutas_global: list[str] = []

    # ── Controles reutilizables ────────────────────────────────

    # Área de log
    log_view = ft.ListView(
        expand=True,
        spacing=2,
        padding=10,
        auto_scroll=True,
    )

    def log(msg: str, color: str = "#cccccc"):
        log_view.controls.append(
            ft.Text(msg, color=color, size=12, font_family="monospace", selectable=True)
        )
        page.update()

    def log_ok(msg):
        log(f"✅  {msg}", "#4ade80")

    def log_warn(msg):
        log(f"⚠️  {msg}", "#fbbf24")

    def log_err(msg):
        log(f"❌  {msg}", "#f87171")

    def log_info(msg):
        log(f"   {msg}", "#94a3b8")

    # Barra de progreso
    progress = ft.ProgressBar(
        value=0,
        color="#4ade80",
        bgcolor="#334155",
        height=6,
    )

    # Etiqueta de estado
    status_text = ft.Text("Listo.", size=11, color="#64748b", italic=True)

    def set_status(msg: str):
        status_text.value = msg
        page.update()

    # Botones (refs para poder habilitarlos/deshabilitarlos)
    btn_ejecutar = ft.Button(
        "▶  Ejecutar extracción",
        icon=ft.Icons.PLAY_ARROW_ROUNDED,
        disabled=True,
        style=ft.ButtonStyle(
            bgcolor={"": "#16a34a", "disabled": "#374151"},
            color={"": "white", "disabled": "#6b7280"},
            padding=20,
        ),
    )

    btn_seleccionar = ft.Button(
        "📁  Seleccionar PDFs",
        icon=ft.Icons.FOLDER_OPEN_ROUNDED,
        style=ft.ButtonStyle(
            bgcolor={"": "#0369a1"},
            color={"": "white"},
            padding=20,
        ),
    )

    btn_limpiar = ft.TextButton(
        "🗑  Limpiar log",
        style=ft.ButtonStyle(color={"": "#94a3b8"}),
    )

    btn_cerrar = ft.TextButton(
        "✕  Cerrar",
        style=ft.ButtonStyle(color={"": "#f87171"}),
    )

    # ── FilePicker ─────────────────────────────────────────────
    file_picker = ft.FilePicker()
    page.overlay.append(file_picker)

    async def on_archivos_seleccionados(e):
        nonlocal rutas_global
        if e.files:
            rutas_global = [f.path for f in e.files]
            log_view.controls.clear()
            log_ok(f"{len(rutas_global)} archivo(s) seleccionado(s):")
            for r in rutas_global:
                log_info(r.split("/")[-1] if "/" in r else r.split("\\")[-1])
            btn_ejecutar.disabled = False
            set_status(f"{len(rutas_global)} PDF(s) listo(s).")
        else:
            set_status("No se seleccionaron archivos.")
        page.update()

    file_picker.on_result = on_archivos_seleccionados

    # ── Lógica de extracción ───────────────────────────────────
    def extraer_convertir():
        try:
            datos, datos_2 = [], []

            for idx, ruta in enumerate(rutas_global, 1):
                nombre = ruta.split("/")[-1] if "/" in ruta else ruta.split("\\")[-1]
                log(f"\n🔍  [{idx}/{len(rutas_global)}]  {nombre}", "#e2e8f0")
                set_status(f"Procesando {idx}/{len(rutas_global)}: {nombre}")
                progress.value = idx / len(rutas_global)
                page.update()

                try:
                    documento = PdfReader(ruta)
                    texto = "\n".join(p.extract_text() or "" for p in documento.pages)
                except Exception as ex:
                    log_err(f"No se pudo leer: {ex}")
                    texto = ""

                # ── Extracciones con regex (idénticas al original) ──────
                fecha = re.search(
                    r"Fecha\s*:\s*("
                    r"\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}"
                    r"\d{1,2}\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}"
                    r"|\d{1,2}[/-]\d{1,2}[/-]\d{4}"
                    r"|\d{4}[/-]\d{1,2}[/-]\d{1,2}"
                    r")",
                    texto,
                    re.IGNORECASE,
                )

                hora_resena = re.search(
                    r"(breve\s+reseña|reseña)\s*[:\-]?\s*(.*?)\b(?:siendo\s+las|a\s+las)\s*(\d{1,2}:\d{2})",
                    texto,
                    re.IGNORECASE | re.DOTALL,
                )
                hora_general = re.search(
                    r"(hora|horario)\s*[:\-]?\s*(\d{1,2}:\d{2})", texto, re.IGNORECASE
                )
                if hora_resena:
                    hora_final = hora_resena.group(3)
                elif hora_general:
                    hora_final = hora_general.group(2)
                else:
                    hora_final = "S/D"

                causa = re.search(
                    r"^\s*(?:causa|hecho)\b\s*[\s*:\-–—]?\s*(.+?)\s*$",
                    texto,
                    re.IGNORECASE | re.MULTILINE,
                )
                jefe = re.search(
                    r"j\s*e\s*f\s*e\s*de\s*servicio:\s*([^\n]+)", texto, re.IGNORECASE
                )
                oficial = re.search(
                    r"(?:Oficial\s*+de\s*+Serv[ií]cio:\s*|Oficial\s*+de\s*+Serv[ií]cio\s*)([^\n]+)",
                    texto,
                    re.IGNORECASE,
                )
                operador = re.search(
                    r"(?:Operador\s*de\s*C[aá]mara|Operador\s*C[aá]mara|Op(?:erador)?\s*C[aá]mara"
                    r"|Aux(?:iliar)?\s+Operador\s+de\s+C[aá]mara|Aux\s*)\s*:?\s*([^\n]+)",
                    texto,
                    re.IGNORECASE,
                )
                sae = re.search(
                    r"(?:SAE\s*Nro.\s*|SAE:\s*°\s*|SAE\s*nº\s*|sae\s*|suceso\s*:\s*|suceso\s*"
                    r"|cad\s*:\s*|SAECAD\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*"
                    r"|sae\s*+nro.|carta\s*:|carta\s*n\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*"
                    r"|SAE\s*N.\s*º\s*|N°\s*|N°\s*:\s*)\s*(\d{8})",
                    texto,
                    re.IGNORECASE,
                )
                resultado = re.search(
                    r"\bresultado\b\s*[:\-–—]?\s*([^\n]+)", texto, re.IGNORECASE
                )
                resultado_texto = resultado.group(1).strip() if resultado else "S/D"

                if operador:
                    op_texto = operador.group(1).strip()
                    operador_valor = (
                        "S/D"
                        if re.match(
                            r"^(breve|rese[nñ]a|resultado|fecha|hora)\b",
                            op_texto,
                            re.IGNORECASE,
                        )
                        else op_texto
                    )
                else:
                    operador_valor = "S/D"

                mapa_numeros = {
                    "un": 1,
                    "uno": 1,
                    "una": 1,
                    "dos": 2,
                    "tres": 3,
                    "cuatro": 4,
                    "cinco": 5,
                    "seis": 6,
                    "siete": 7,
                    "ocho": 8,
                    "nueve": 9,
                    "diez": 10,
                }

                def extraer_numero_resultado(t):
                    t = t.lower()
                    for pat in (r"\b(\d+)\s*\(", r"\b(\d+)\b", r"\((\d+)\)"):
                        m = re.search(pat, t)
                        if m:
                            return int(m.group(1))
                    for palabra, valor in mapa_numeros.items():
                        if re.search(rf"\b{palabra}\b", t):
                            return valor
                    return ""

                detenido_valor = contraventor_valor = ""
                if re.search(r"\bdeten", resultado_texto, re.IGNORECASE):
                    detenido_valor = extraer_numero_resultado(resultado_texto)
                if re.search(r"\bcontra", resultado_texto, re.IGNORECASE):
                    contraventor_valor = extraer_numero_resultado(resultado_texto)
                resultado_final = (
                    "" if (detenido_valor or contraventor_valor) else resultado_texto
                )

                fecha_match = re.search(
                    r"Fecha\s*:\s*(.+?)(?:\s+Hora\b|\n|$)", texto, re.IGNORECASE
                )
                if fecha_match:
                    fecha_texto_raw = fecha_match.group(1).strip()
                else:
                    fb = re.search(
                        r"(\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+(?:\s+de(?:l)?\s+\d{2,4})?|\d{1,2}[\/\-]\d{1,2}(?:[\/\-]\d{2,4})?|\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2})",
                        texto,
                        re.IGNORECASE,
                    )
                    fecha_texto_raw = fb.group(1).strip() if fb else "S/D"
                fecha_formateada = formatear_fecha(fecha_texto_raw)

                datos.append(
                    {
                        "RUTA": ruta,
                        "FECHA": fecha_formateada,
                        "TIPIFICACION": "",
                        "CAUSA": causa.group(1).strip() if causa else "S/D",
                        "SAE": int(sae.group(1).strip()) if sae else "S/D",
                        "HORA": hora_final,
                        "OPERADOR DE CAMARA": operador_valor,
                        "OFICIAL DE SERVICIO": oficial.group(1).strip()
                        if oficial
                        else "S/D",
                        "CONTRAVENTOR": contraventor_valor,
                        "DETENIDO": detenido_valor,
                        "JEFE DE SERVICIO": jefe.group(1).strip() if jefe else "S/D",
                        "RESULTADO": resultado_final,
                    }
                )

                resena_bloque = re.search(
                    r"(?:rese[nñ]a)\s*:?\s*(.*?)Resultado\s*:",
                    texto,
                    re.IGNORECASE | re.DOTALL,
                )
                resena_limpia = (
                    re.sub(r"\s+", " ", resena_bloque.group(1)).strip()
                    if resena_bloque
                    else "S/D"
                )
                datos_2.append(
                    f"Fecha: {fecha_formateada}\n"
                    f"Reseña: {resena_limpia}\n"
                    f"Resultado: {resultado.group(1).strip() if resultado else 'S/D'}"
                )
                log_ok(
                    f"Extraído — Fecha: {fecha_formateada}  SAE: {sae.group(1) if sae else 'S/D'}"
                )

            # ── Guardar Excel ───────────────────────────────────────
            log("\n💾  Guardando Excel…", "#e2e8f0")
            try:
                df_existente = pd.read_excel("contenido_archivos.xlsx")
                df_final = pd.concat(
                    [df_existente, pd.DataFrame(datos)], ignore_index=True
                )
            except FileNotFoundError:
                df_final = pd.DataFrame(datos)
            df_final.to_excel("contenido_archivos.xlsx", index=False)

            try:
                wb = openpyxl.load_workbook("contenido_archivos.xlsx")
                ws = wb.active

                palabras_limpiar = [
                    "personal",
                    "contratado",
                    "auxiliar",
                    "pers",
                    "contr",
                    "aux",
                    "cont",
                    "oficial",
                    "of",
                    "operador",
                    "oper",
                    "lp",
                    "primero",
                    "mayor",
                    "inspector",
                    "subrio",
                    "subcomisario",
                ]
                patrones = [
                    re.compile(r"\b" + re.escape(p) + r"\.?\b", flags=re.IGNORECASE)
                    for p in palabras_limpiar
                ]

                for col in ["G", "H", "K"]:
                    for celda in ws[col]:
                        if celda.row == 1 or not celda.value:
                            continue
                        t = str(celda.value)
                        for pat in patrones:
                            t = pat.sub("", t)
                        t = re.sub(r"[.\-:;]", " ", t)
                        t = re.sub(r"\d+", " ", t)
                        t = re.sub(r"\s+", " ", t).strip()
                        celda.value = t or "S/D"

                for col in ["B", "E", "F", "I", "J"]:
                    for celda in ws[col]:
                        celda.alignment = Alignment(horizontal="center")

                for row in range(2, ws.max_row + 1):
                    celda = ws.cell(row=row, column=2)
                    if (
                        celda.value
                        and celda.value != "S/D"
                        and isinstance(celda.value, str)
                    ):
                        try:
                            celda.value = datetime.strptime(celda.value, "%d/%m/%Y")
                            celda.number_format = "DD/MM/AAAA"
                        except Exception:
                            pass

                for row in range(2, ws.max_row + 1):
                    ruta_pdf = ws.cell(row=row, column=1).value
                    if (
                        ruta_pdf
                        and isinstance(ruta_pdf, str)
                        and ruta_pdf.lower().endswith(".pdf")
                    ):
                        ws.cell(row=row, column=1).hyperlink = Hyperlink(
                            ref=f"A{row}", target=str(ruta_pdf)
                        )
                        ws.cell(row=row, column=1).font = Font(
                            color="0000FF", underline="single"
                        )

                wb.save("contenido_archivos.xlsx")
            except Exception as ex:
                log_warn(f"Error al formatear Excel: {ex}")

            # ── Guardar Word ────────────────────────────────────────
            log("💾  Guardando Word…", "#e2e8f0")
            try:
                doc = Document("reseña_archivos.docx")
            except Exception:
                doc = Document()
                encabezado = doc.add_heading("División CENTROS DE MONITOREO URBANO", 0)
                encabezado.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                encabezado.runs[0].font.size = Pt(20)

                table = doc.add_table(rows=10, cols=5)
                table.style = "Table Grid"
                fila0 = table.rows[0]
                fila0.cells[0].merge(fila0.cells[4])
                table.cell(0, 0).text = "PARTE SEMANAL DEL XX/XX/2025 AL XX/XX/2025"
                table.cell(1, 0).text = "MES DE DICIEMBRE/DIA"
                table.cell(1, 1).text = "EVENTOS DEL PERÍODO"
                table.cell(1, 2).text = "EVENTOS POSITIVOS"
                table.cell(1, 3).text = "IMPUTADOS"
                table.cell(1, 4).text = "EVENTOS RELEVANTES"

                celda_p = table.cell(0, 0).paragraphs[0]
                celda_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = celda_p.runs[0]
                run.font.size = Pt(14)
                run.font.name = "Arial"
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                fondo = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls("w")))
                table.rows[0].cells[0]._tc.get_or_add_tcPr().append(fondo)

                for i in range(5):
                    p = table.cell(1, i).paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.runs[0]
                    r.font.size = Pt(10)
                    r.font.name = "Arial"
                    r.font.bold = True

                res_enc = doc.add_heading("\n\nReseñas extraídas", 0)
                res_enc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                res_enc.runs[0].font.size = Pt(14)

            for item in datos_2:
                doc.add_paragraph(item)
                doc.add_paragraph(
                    "--------------------------------------------------------------"
                )
            doc.save("reseña_archivos.docx")

            progress.value = 1.0
            log_ok("Proceso completado.")
            set_status("✅ Guardado en contenido_archivos.xlsx y reseña_archivos.docx")
            page.open(
                ft.SnackBar(
                    content=ft.Text(
                        "✅ Datos guardados en contenido_archivos.xlsx y reseña_archivos.docx"
                    ),
                    bgcolor="#16a34a",
                )
            )

        except Exception as ex:
            traceback.print_exc()
            log_err(str(ex))
            set_status("Error durante el procesamiento.")
            page.open(
                ft.SnackBar(
                    content=ft.Text(f"❌ Error: {ex}"),
                    bgcolor="#dc2626",
                )
            )

        finally:
            btn_ejecutar.disabled = False
            btn_seleccionar.disabled = False
            page.update()

    # ── Callbacks de botones ───────────────────────────────────
    async def click_seleccionar(_):
        await file_picker.pick_files(
            dialog_title="Seleccionar archivos PDF",
            allowed_extensions=["pdf"],
            allow_multiple=True,
        )

    def click_ejecutar(_):
        if not rutas_global:
            page.open(
                ft.SnackBar(
                    content=ft.Text("⚠️  Seleccioná archivos PDF primero."),
                    bgcolor="#b45309",
                )
            )
            page.update()
            return
        btn_ejecutar.disabled = True
        btn_seleccionar.disabled = True
        progress.value = None  # modo indeterminado
        page.update()
        threading.Thread(target=extraer_convertir, daemon=True).start()

    def click_limpiar(_):
        log_view.controls.clear()
        log("Log limpiado.", "#64748b")
        progress.value = 0
        page.update()

    btn_seleccionar.on_click = click_seleccionar
    btn_ejecutar.on_click = click_ejecutar
    btn_limpiar.on_click = click_limpiar
    btn_cerrar.on_click = lambda _: page.window.close()

    # ── Layout ─────────────────────────────────────────────────
    header = ft.Container(
        content=ft.Row(
            [
                ft.Icon(ft.Icons.PICTURE_AS_PDF_ROUNDED, color="#f87171", size=28),
                ft.Column(
                    [
                        ft.Text(
                            "Extractor PDF → Excel / Word",
                            size=17,
                            weight=ft.FontWeight.BOLD,
                            color="white",
                        ),
                        ft.Text(
                            "División Centros de Monitoreo Urbano",
                            size=10,
                            color="#64748b",
                        ),
                    ],
                    spacing=0,
                ),
            ],
            spacing=12,
        ),
        padding=ft.padding.symmetric(horizontal=16, vertical=12),
        bgcolor="#0f172a",
    )

    log_container = ft.Container(
        content=log_view,
        expand=True,
        bgcolor="#0f172a",
        border=ft.border.all(1, "#1e293b"),
        border_radius=8,
        margin=ft.margin.symmetric(horizontal=12),
    )

    buttons_row = ft.Container(
        content=ft.Row(
            [
                btn_seleccionar,
                btn_ejecutar,
                btn_limpiar,
                ft.Container(expand=True),
                btn_cerrar,
            ],
            spacing=8,
        ),
        padding=ft.padding.symmetric(horizontal=12, vertical=8),
    )

    status_bar = ft.Container(
        content=ft.Row(
            [
                ft.Icon(ft.Icons.INFO_OUTLINE, size=13, color="#475569"),
                status_text,
            ],
            spacing=6,
        ),
        padding=ft.padding.symmetric(horizontal=14, vertical=4),
        bgcolor="#0f172a",
    )

    page.add(
        ft.Column(
            [
                header,
                ft.Divider(height=1, color="#1e293b"),
                log_container,
                ft.Container(
                    content=progress,
                    padding=ft.padding.symmetric(horizontal=12, vertical=4),
                ),
                buttons_row,
                ft.Divider(height=1, color="#1e293b"),
                status_bar,
            ],
            expand=True,
            spacing=0,
        )
    )

    # Mensaje inicial
    log("Hacé clic en '📁 Seleccionar PDFs' para comenzar.", "#64748b")
    page.update()


# ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    ft.app(target=main)
