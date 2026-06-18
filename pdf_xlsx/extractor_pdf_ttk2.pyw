import re
import threading
import traceback
from datetime import datetime

import pandas as pd
import openpyxl
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Alignment
from openpyxl.worksheet.hyperlink import Hyperlink

#import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as ttk
#from ttkbootstrap.constants import *
try:
    from ttkbootstrap.widgets.scrolled import ScrolledText
except ImportError:
    from ttkbootstrap.scrolled import ScrolledText

# ──────────────────────────────────────────────
#  Estado global
# ──────────────────────────────────────────────
rutas_global = []


# ──────────────────────────────────────────────
#  Helpers
# ──────────────────────────────────────────────
def formatear_fecha(fecha_texto):
    """Convierte distintos formatos de fecha y devuelve solo 'dd/mm'."""
    if not fecha_texto or fecha_texto == "S/D":
        return "S/D"
    try:
        fecha_limpia = fecha_texto.strip()
        meses_espanol = {
            "enero": 1, "ene": 1, "febrero": 2, "feb": 2,
            "marzo": 3, "mar": 3, "abril": 4, "abr": 4,
            "mayo": 5, "may": 5, "junio": 6, "jun": 6,
            "julio": 7, "jul": 7, "agosto": 8, "ago": 8,
            "septiembre": 9, "sep": 9, "setiembre": 9,
            "octubre": 10, "oct": 10, "noviembre": 11, "nov": 11,
            "diciembre": 12, "dic": 12,
        }

        m = re.search(r"(\d{1,2})\s+de\s+([A-záéíóúñ]+)", fecha_limpia, re.IGNORECASE)
        if not m:
            m = re.search(r"(\d{1,2})\s+([A-záéíóúñ]+)", fecha_limpia, re.IGNORECASE)
        if m:
            dia = int(m.group(1))
            mes = meses_espanol.get(m.group(2).lower().replace(".", ""))
            if mes:
                return f"{dia:02d}/{mes:02d}"

        m = re.search(r"(\d{1,2})[\/\-](\d{1,2})(?:[\/\-]\d{2,4})?", fecha_limpia)
        if m:
            return f"{int(m.group(1)):02d}/{int(m.group(2)):02d}"

        m = re.search(r"(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})", fecha_limpia)
        if m:
            return f"{int(m.group(3)):02d}/{int(m.group(2)):02d}"

        for fmt in ("%d/%m/%Y", "%d/%m", "%Y-%m-%d", "%m-%d"):
            try:
                return datetime.strptime(fecha_limpia, fmt).strftime("%d/%m")
            except Exception:
                continue

        return fecha_limpia
    except Exception as e:
        print(f"Error al formatear fecha '{fecha_texto}': {e}")
        return fecha_texto


def log(msg: str):
    """Escribe una línea en el área de log de la UI."""
    txt_log.configure(state="normal")
    txt_log.insert("end", msg + "\n")
    txt_log.see("end")
    txt_log.configure(state="disabled")


def set_status(msg: str):
    """Actualiza la etiqueta de estado en la barra inferior."""
    lbl_status.config(text=msg)


# ──────────────────────────────────────────────
#  Acciones de UI
# ──────────────────────────────────────────────
def seleccionar_archivos():
    global rutas_global
    try:
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")],
        )
        if archivos:
            rutas_global = list(archivos)
            txt_log.configure(state="normal")
            txt_log.delete("1.0", "end")
            txt_log.configure(state="disabled")
            log(f"✅ {len(rutas_global)} archivo(s) seleccionado(s):\n")
            for r in rutas_global:
                log(f"   📄 {r.split('/')[-1]}")
            set_status(f"{len(rutas_global)} PDF(s) listos para procesar.")
            btn_ejecutar.config(state="normal")
        else:
            set_status("No se seleccionaron archivos.")
    except Exception as e:
        traceback.print_exc()
        messagebox.showerror("Error", f"Error al seleccionar archivos:\n{e}")


def iniciar_extraccion():
    """Lanza el procesamiento en un hilo separado para no bloquear la UI."""
    btn_ejecutar.config(state="disabled")
    btn_seleccionar.config(state="disabled")
    progress_bar.config(mode="indeterminate")
    progress_bar.start(10)
    set_status("Procesando PDFs…")
    hilo = threading.Thread(target=extraer_convertir, daemon=True)
    hilo.start()


def finalizar_ui(exito: bool, mensaje: str):
    """Restablece la UI tras finalizar el procesamiento (llamado desde el hilo)."""
    progress_bar.stop()
    progress_bar.config(mode="determinate", value=100 if exito else 0)
    btn_seleccionar.config(state="normal")
    btn_ejecutar.config(state="normal")
    set_status(mensaje)
    if exito:
        messagebox.showinfo("Éxito", mensaje)
    else:
        messagebox.showerror("Error", mensaje)


# ──────────────────────────────────────────────
#  Lógica de extracción (sin cambios funcionales)
# ──────────────────────────────────────────────
def extraer_convertir():
    try:
        if not rutas_global:
            messagebox.showwarning("Advertencia", "Seleccioná archivos PDF primero.")
            finalizar_ui(False, "Sin archivos seleccionados.")
            return

        datos, datos_2 = [], []
        error = ""

        for idx, ruta in enumerate(rutas_global, 1):
            nombre = ruta.split("/")[-1]
            log(f"\n🔍 [{idx}/{len(rutas_global)}] Procesando: {nombre}")
            set_status(f"Procesando {idx}/{len(rutas_global)}: {nombre}")

            try:
                documento = PdfReader(ruta)
                texto = "\n".join(
                    p.extract_text() or "" for p in documento.pages
                )
            except Exception as e:
                error += f"Error al leer {ruta}: {e}\n"
                log("   ⚠️  No se pudo leer el archivo.")
                texto = ""

            fecha = re.search(
                r"Fecha\s*:\s*("
                r"\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}"
                r"\d{1,2}\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+\s+de(?:l)?\s+\d{4}"
                r"|\d{1,2}[/-]\d{1,2}[/-]\d{4}"
                r"|\d{4}[/-]\d{1,2}[/-]\d{1,2}"
                r")",
                texto, re.IGNORECASE,
            )

            hora_resena = re.search(
                r"(breve\s+reseña|reseña)\s*[:\-]?\s*(.*?)\b(?:siendo\s+las|a\s+las)\s*(\d{1,2}:\d{2})",
                texto, re.IGNORECASE | re.DOTALL,
            )
            hora_general = re.search(
                r"(hora|horario)\s*[:\-]?\s*(\d{1,2}:\d{2})", texto, re.IGNORECASE
            )
            if hora_resena:
                hora_final = hora_resena.group(3)
            elif hora_general:
                hora_final = hora_general.group(2)
            else:
                hora_final = "S/D"

            causa = re.search(
                r"^\s*(?:causa|hecho)\b\s*[\s*:\-–—]?\s*(.+?)\s*$",
                texto, re.IGNORECASE | re.MULTILINE,
            )
            jefe = re.search(r"j\s*e\s*f\s*e\s*de\s*servicio:\s*([^\n]+)", texto, re.IGNORECASE)
            oficial = re.search(
                r"(?:Oficial\s*+de\s*+Serv[ií]cio:\s*|Oficial\s*+de\s*+Serv[ií]cio\s*)([^\n]+)",
                texto, re.IGNORECASE,
            )
            operador = re.search(
                r"(?:Operador\s*de\s*C[aá]mara|Operador\s*C[aá]mara|Op(?:erador)?\s*C[aá]mara"
                r"|Aux(?:iliar)?\s+Operador\s+de\s+C[aá]mara|Aux\s*)\s*:?\s*([^\n]+)",
                texto, re.IGNORECASE,
            )
            sae = re.search(
                r"(?:SAE\s*Nro.\s*|SAE:\s*°\s*|SAE\s*nº\s*|sae\s*|suceso\s*:\s*|suceso\s*"
                r"|cad\s*:\s*|SAECAD\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*"
                r"|sae\s*+nro.|carta\s*:|carta\s*n\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*"
                r"|SAE\s*N.\s*º\s*|N°\s*|N°\s*:\s*)\s*(\d{8})",
                texto, re.IGNORECASE,
            )
            resultado = re.search(r"\bresultado\b\s*[:\-–—]?\s*([^\n]+)", texto, re.IGNORECASE)
            resultado_texto = resultado.group(1).strip() if resultado else "S/D"

            if operador:
                op_texto = operador.group(1).strip()
                operador_valor = (
                    "S/D"
                    if re.match(r"^(breve|rese[nñ]a|resultado|fecha|hora)\b", op_texto, re.IGNORECASE)
                    else op_texto
                )
            else:
                operador_valor = "S/D"

            mapa_numeros = {
                "un": 1, "uno": 1, "una": 1, "dos": 2, "tres": 3,
                "cuatro": 4, "cinco": 5, "seis": 6, "siete": 7,
                "ocho": 8, "nueve": 9, "diez": 10,
            }

            def extraer_numero_resultado(t):
                t = t.lower()
                for pat in (r"\b(\d+)\s*\(", r"\b(\d+)\b", r"\((\d+)\)"):
                    m = re.search(pat, t)
                    if m:
                        return int(m.group(1))
                for palabra, valor in mapa_numeros.items():
                    if re.search(rf"\b{palabra}\b", t):
                        return valor
                return ""

            detenido_valor = contraventor_valor = ""
            if re.search(r"\bdeten", resultado_texto, re.IGNORECASE):
                detenido_valor = extraer_numero_resultado(resultado_texto)
            if re.search(r"\bcontra", resultado_texto, re.IGNORECASE):
                contraventor_valor = extraer_numero_resultado(resultado_texto)
            resultado_final = "" if (detenido_valor or contraventor_valor) else resultado_texto

            fecha_match = re.search(r"Fecha\s*:\s*(.+?)(?:\s+Hora\b|\n|$)", texto, re.IGNORECASE)
            if fecha_match:
                fecha_texto = fecha_match.group(1).strip()
            else:
                fb = re.search(
                    r"(\d{1,2}\s+de\s+[A-Za-zÁÉÍÓÚáéíóúñÑ]+(?:\s+de(?:l)?\s+\d{2,4})?|\d{1,2}[\/\-]\d{1,2}(?:[\/\-]\d{2,4})?|\d{4}[\/\-]\d{1,2}[\/\-]\d{1,2})",
                    texto, re.IGNORECASE,
                )
                fecha_texto = fb.group(1).strip() if fb else "S/D"
            fecha_formateada = formatear_fecha(fecha_texto)

            datos.append({
                "RUTA": ruta,
                "FECHA": fecha_formateada,
                "TIPIFICACION": "",
                "CAUSA": causa.group(1).strip() if causa else "S/D",
                "SAE": int(sae.group(1).strip()) if sae else "S/D",
                "HORA": hora_final,
                "OPERADOR DE CAMARA": operador_valor,
                "OFICIAL DE SERVICIO": oficial.group(1).strip() if oficial else "S/D",
                "CONTRAVENTOR": contraventor_valor,
                "DETENIDO": detenido_valor,
                "JEFE DE SERVICIO": jefe.group(1).strip() if jefe else "S/D",
                "RESULTADO": resultado_final,
            })

            resena_bloque = re.search(
                r"(?:rese[nñ]a)\s*:?\s*(.*?)Resultado\s*:", texto, re.IGNORECASE | re.DOTALL
            )
            resena_limpia = (
                re.sub(r"\s+", " ", resena_bloque.group(1)).strip()
                if resena_bloque else "S/D"
            )
            datos_2.append(
                f"Fecha: {fecha_formateada}\n"
                f"Reseña: {resena_limpia}\n"
                f"Resultado: {resultado.group(1).strip() if resultado else 'S/D'}"
            )
            log(f"   ✔ Extraído — Fecha: {fecha_formateada}  SAE: {sae.group(1) if sae else 'S/D'}")

        # ── Excel ──────────────────────────────
        log("\n💾 Guardando Excel…")
        try:
            df_existente = pd.read_excel("contenido_archivos.xlsx")
            df_final = pd.concat([df_existente, pd.DataFrame(datos)], ignore_index=True)
        except FileNotFoundError:
            df_final = pd.DataFrame(datos)
        df_final.to_excel("contenido_archivos.xlsx", index=False)

        try:
            wb = openpyxl.load_workbook("contenido_archivos.xlsx")
            ws = wb.active

            palabras_limpiar = [
                "personal", "contratado", "auxiliar", "pers", "contr", "aux",
                "cont", "oficial", "of", "operador", "oper", "lp", "primero",
                "mayor", "inspector", "subrio", "subcomisario",
            ]
            patrones = [
                re.compile(r"\b" + re.escape(p) + r"\.?\b", flags=re.IGNORECASE)
                for p in palabras_limpiar
            ]

            for col in ["G", "H", "K"]:
                for celda in ws[col]:
                    if celda.row == 1 or not celda.value:
                        continue
                    t = str(celda.value)
                    for pat in patrones:
                        t = pat.sub("", t)
                    t = re.sub(r"[.\-:;]", " ", t)
                    t = re.sub(r"\d+", " ", t)
                    t = re.sub(r"\s+", " ", t).strip()
                    celda.value = t or "S/D"

            for col in ["B", "E", "F", "I", "J"]:
                for celda in ws[col]:
                    celda.alignment = Alignment(horizontal="center")

            for row in range(2, ws.max_row + 1):
                celda = ws.cell(row=row, column=2)
                if celda.value and celda.value != "S/D" and isinstance(celda.value, str):
                    try:
                        celda.value = datetime.strptime(celda.value, "%d/%m/%Y")
                        celda.number_format = "DD/MM/AAAA"
                    except Exception:
                        pass

            for row in range(2, ws.max_row + 1):
                ruta_pdf = ws.cell(row=row, column=1).value
                if ruta_pdf and isinstance(ruta_pdf, str) and ruta_pdf.lower().endswith(".pdf"):
                    ws.cell(row=row, column=1).hyperlink = Hyperlink(ref=f"A{row}", target=str(ruta_pdf))
                    ws.cell(row=row, column=1).font = Font(color="0000FF", underline="single")

            wb.save("contenido_archivos.xlsx")
        except Exception as e:
            log(f"   ⚠️  Error al formatear Excel: {e}")

        # ── Word ───────────────────────────────
        log("💾 Guardando Word…")
        try:
            doc = Document("reseña_archivos.docx")
        except Exception:
            doc = Document()
            encabezado = doc.add_heading("División CENTROS DE MONITOREO URBANO", 0)
            encabezado.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            encabezado.runs[0].font.size = Pt(20)

            table = doc.add_table(rows=10, cols=5)
            table.style = "Table Grid"
            row0 = table.rows[0]
            row0.cells[0].merge(row0.cells[4])
            table.cell(0, 0).text = "PARTE SEMANAL DEL XX/XX/2025 AL XX/XX/2025"
            table.cell(1, 0).text = "MES DE DICIEMBRE/DIA"
            table.cell(1, 1).text = "EVENTOS DEL PERÍODO"
            table.cell(1, 2).text = "EVENTOS POSITIVOS"
            table.cell(1, 3).text = "IMPUTADOS"
            table.cell(1, 4).text = "EVENTOS RELEVANTES"

            celda_p = table.cell(0, 0).paragraphs[0]
            celda_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = celda_p.runs[0]
            run.font.size = Pt(14)
            run.font.name = "Arial"
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            fondo = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls("w")))
            table.rows[0].cells[0]._tc.get_or_add_tcPr().append(fondo)

            for i in range(5):
                p = table.cell(1, i).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.runs[0]
                r.font.size = Pt(10)
                r.font.name = "Arial"
                r.font.bold = True

            res_enc = doc.add_heading("\n\nReseñas extraídas", 0)
            res_enc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            res_enc.runs[0].font.size = Pt(14)

        for item in datos_2:
            doc.add_paragraph(item)
            doc.add_paragraph("--------------------------------------------------------------")
        doc.save("reseña_archivos.docx")

        log("\n✅ Proceso completado con éxito.")
        root.after(0, finalizar_ui, True,
                   "Datos guardados en contenido_archivos.xlsx y reseña_archivos.docx")

    except Exception as e:
        traceback.print_exc()
        log(f"\n❌ Error: {e}")
        root.after(0, finalizar_ui, False, f"Error durante el procesamiento:\n{e}")


# ──────────────────────────────────────────────
#  Interfaz gráfica con ttkbootstrap
# ──────────────────────────────────────────────
root = ttk.Window(
    title="Extractor PDF → Excel / Word",
    themename="darkly",       # Temas disponibles: cosmo, flatly, litera, minty, lumen,
                               #   sandstone, yeti, pulse, united, morph, journal,
                               #   darkly, superhero, solar, cyborg, vapor, simplex, cerculean
    size=(860, 520),
    resizable=(True, True),
)
root.place_window_center()

# ── Encabezado ────────────────────────────────
frm_header = ttk.Frame(root, padding=(12, 10))
frm_header.pack(fill="x")

ttk.Label(
    frm_header,
    text="📋  Extractor de datos PDF",
    font=("Helvetica", 16, "bold"),
    bootstyle="inverse-dark",
).pack(side="left")

ttk.Label(
    frm_header,
    text="División Centros de Monitoreo Urbano",
    font=("Helvetica", 9),
    bootstyle="secondary",
).pack(side="left", padx=14, pady=(4, 0))

# ── Área de log (ScrolledText) ────────────────
frm_log = ttk.LabelFrame(root, text=" Registro de procesamiento ", bootstyle="secondary")
frm_log.pack(fill="both", expand=True, padx=12, pady=(0, 6))

txt_log = ScrolledText(frm_log, height=14, autohide=True, state="disabled",
                       font=("Consolas", 9), padding=6)
txt_log.pack(fill="both", expand=True, padx=4, pady=4)

log("Haz clic en 'Seleccionar Archivos PDF' para comenzar.")

# ── Barra de progreso ─────────────────────────
progress_bar = ttk.Progressbar(root, bootstyle="success-striped", length=300)
progress_bar.pack(fill="x", padx=12, pady=(0, 4))

# ── Botones ───────────────────────────────────
frm_btns = ttk.Frame(root, padding=(12, 4))
frm_btns.pack(fill="x")

btn_seleccionar = ttk.Button(
    frm_btns,
    text="📁  Seleccionar Archivos PDF",
    bootstyle="info-outline",
    width=26,
    command=seleccionar_archivos,
)
btn_seleccionar.pack(side="left", padx=(0, 8))

btn_ejecutar = ttk.Button(
    frm_btns,
    text="▶  Ejecutar extracción",
    bootstyle="success",
    width=22,
    command=iniciar_extraccion,
    state="disabled",         # Se habilita solo cuando hay archivos seleccionados
)
btn_ejecutar.pack(side="left", padx=(0, 8))

btn_limpiar = ttk.Button(
    frm_btns,
    text="🗑  Limpiar log",
    bootstyle="secondary-outline",
    width=14,
    command=lambda: (
        txt_log.configure(state="normal"),
        txt_log.delete("1.0", "end"),
        txt_log.configure(state="disabled"),
        log("Log limpiado."),
    ),
)
btn_limpiar.pack(side="left")

btn_cerrar = ttk.Button(
    frm_btns,
    text="✕  Cerrar",
    bootstyle="danger-outline",
    width=12,
    command=root.quit,
)
btn_cerrar.pack(side="right")

# ── Barra de estado ───────────────────────────
frm_status = ttk.Frame(root, padding=(12, 3))
frm_status.pack(fill="x", side="bottom")
ttk.Separator(root).pack(fill="x")

lbl_status = ttk.Label(frm_status, text="Listo.", bootstyle="secondary", font=("Helvetica", 8))
lbl_status.pack(side="left")

# ──────────────────────────────────────────────
if __name__ == "__main__":
    root.mainloop()
