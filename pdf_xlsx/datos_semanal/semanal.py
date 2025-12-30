import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_reporte_desde_excel(archivo_excel, columna_fecha, parametros_a_contar, fecha_inicio, fecha_fin, nombre_word="reporte.docx"):
    

    # Convertir strings a fecha
    fecha_inicio = pd.to_datetime(fecha_inicio)
    fecha_fin = pd.to_datetime(fecha_fin)
    
#TODO  es una prueba para ver el funcionamiento de la extension TODO

    # ======== LEER EXCEL ==========
    df = pd.read_excel(archivo_excel)

    # Asegurar que la columna de fecha es datetime
    df[columna_fecha] = pd.to_datetime(df[columna_fecha], errors='coerce')

    # ======== FILTRAR ==========
    df_filtrado = df[(df[columna_fecha] >= fecha_inicio) & (df[columna_fecha] <= fecha_fin)]

    if df_filtrado.empty:
        print("No hay datos en el rango indicado.")
        return None

    # ======== AGRUPAR POR DÍA ==========
    # Convertir la fecha a solo día
    df_filtrado["DIA"] = df_filtrado[columna_fecha].dt.date

    # Contar por día los parámetros dados
    conteos = df_filtrado.groupby("DIA")[parametros_a_contar].agg("count").reset_index()
    
    # Contar PC solo cuando TIPIFICACIÓN != "COLABORACION" y PC tenga contenido
    df_filtrado_pc = df_filtrado[df_filtrado["TIPIFICACIÓN"] != "COLABORACION"].copy()
    
    # Contar filas donde PC NO está vacío (cuenta como 1 por horario)
    pc_por_dia = df_filtrado_pc[df_filtrado_pc["PC"].notna() & (df_filtrado_pc["PC"] != "")].groupby("DIA").size().reset_index(name="PC_suma")
    
    # Contar "FINALIZA CON IMPUTADO/S"
    df_filtrado_cierre = df_filtrado[df_filtrado["TIPO DE CIERRE"] == "FINALIZA CON IMPUTADO/S"].copy()
    cierre_por_dia = df_filtrado_cierre.groupby("DIA").size().reset_index(name="FINALIZA_IMPUTADOS")
    
    # Sumar CONTRAVENTORES y DETENIDOS solo cuando TIPO DE CIERRE == "FINALIZA CON IMPUTADO/S"
    imputados_por_dia = df_filtrado_cierre.groupby("DIA")[["CONTRAVENTORES", "DETENIDOS"]].sum().reset_index()
    imputados_por_dia["IMPUTADOS_TOTAL"] = imputados_por_dia["CONTRAVENTORES"] + imputados_por_dia["DETENIDOS"]
    
    # Fusionar los conteos con PC, cierre e imputados
    conteos = conteos.merge(pc_por_dia, on="DIA", how="left")
    conteos = conteos.merge(cierre_por_dia, on="DIA", how="left")
    conteos = conteos.merge(imputados_por_dia[["DIA", "IMPUTADOS_TOTAL"]], on="DIA", how="left")
    conteos["PC_suma"] = conteos["PC_suma"].fillna(0).astype(int)
    conteos["FINALIZA_IMPUTADOS"] = conteos["FINALIZA_IMPUTADOS"].fillna(0).astype(int)
    conteos["IMPUTADOS_TOTAL"] = conteos["IMPUTADOS_TOTAL"].fillna(0).astype(int)

    # ======== CREAR WORD ==========
    doc = Document()

    # Título
    titulo = doc.add_heading("Reporte por Día", level=0)
    titulo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.size = Pt(20)

    # Crear tabla
    columnas = ["DIA"] + parametros_a_contar + ["PC_suma", "IMPUTADOS_TOTAL", "FINALIZA_IMPUTADOS"]
    tabla = doc.add_table(rows=1, cols=len(columnas))

    # Encabezados
    for i, col in enumerate(columnas):
        tabla.cell(0, i).text = col

    # Cargar datos
    for _, fila in conteos.iterrows():
        row = tabla.add_row().cells
        for i, col in enumerate(columnas):
            row[i].text = str(fila[col])

    # Guardar archivo
    doc.save(nombre_word)
    print(f"Reporte generado correctamente: {nombre_word}")

    return conteos


generar_reporte_desde_excel("/Users/alexisjankowicz/Python/pdf_xlsx/datos_semanal/CMU - OCTUBRE.xlsx", "FECHA", ["TIPIFICACIÓN"], "2025/10/01", "2025/10/07", "reporte_semanal.docx")
# Ejemplo de uso:
# generar_reporte_desde_excel("datos.xlsx", "Fecha", ["Parametro1", "Parametro2"], "2024-01-01", "2024-01-31", "reporte_semanal.docx")
