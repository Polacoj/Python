import tkinter as tk
from tkinterdnd2 import DND_FILES, TkinterDnD
from pypdf import PdfReader
import re
import pandas as pd
from docx import Document

# Variables globales para almacenar rutas y contenido
rutas_global = []
contenido_global = ""


def pegar(evento):
    global rutas_global, contenido_global
    archivo = evento.data
    # Extrae rutas entre llaves o rutas simples
    rutas = re.findall(r'\{([^}]+)\}|([^\s]+)', archivo)
    # Normaliza la lista de rutas
    rutas = [ruta[0] if ruta[0] else ruta[1] for ruta in rutas]
    rutas_global = rutas  # Guarda las rutas para usarlas al aceptar
    archivos_texto = "Pegado:\n" + "\n".join(rutas)
    ventana_pegado.config(text=archivos_texto)
    print(f"Archivos pegados: {rutas}")

    contenido = ""
    for ruta in rutas:
        try:
            documento = PdfReader(ruta)
            contenido += "----------" + ruta + "----------------\n"
            for pagina in documento.pages:
                texto_pagina = pagina.extract_text()
                if texto_pagina:
                    contenido += texto_pagina + "\n"
        except Exception as e:
            contenido += f"Error al leer {ruta}: {str(e)}\n"
    contenido_global = contenido  # Guarda el contenido para usarlo al aceptar
    print(contenido)


def aceptar():
    global rutas_global, contenido_global
    # Cambia el texto del label para indicar aceptación
    ventana_pegado.config(text="Archivos cargados")
    print("Archivos cargados")
    # Guarda el contenido en un archivo .txt
    if contenido_global:
        with open("contenido_archivos.txt", "a", encoding="utf-8") as f:
            f.write(contenido_global)
            f.write("\n\n")
        print("Contenido guardado en contenido_archivos.txt")


def extraer_convertir():
    datos, datos_2 = [], []
    with open("contenido_archivos.txt", "r", encoding="utf-8") as f:
        contenido = f.read()

    bloques = re.split(r'-{10,}(.+?)-{10,}\n', contenido)
    for i in range(1, len(bloques), 2):
        ruta = bloques[i].strip()
        texto = bloques[i+1].strip() if i+1 < len(bloques) else ""

        fecha = re.search(r'Fecha:\s*(.*?2025)', texto, re.IGNORECASE)
        hora = re.search(r'(Hora\s*|Horario\s*)[:\- ]+([^\n]+)', texto)
        causa = re.search(
            r'(?:Causa\s*:\s*|hecho\s*:\s*)([^\n]+)', texto, re.IGNORECASE)
        jefe = re.search(
            r'Jefe\s*+de\s*+Servicio:\s*([^\n]+)', texto, re.IGNORECASE)
        oficial = re.search(
            r'Oficial\s*+de\s*+Serv[ií]cio:\s*([^\n]+)', texto, re.IGNORECASE)
        operador = re.search(
            r'Operador\s*+de\s*+C[aá]mara:\s*([^\n]+)', texto, re.IGNORECASE)
        sae = re.search(
            r'(?:SAE\s*nº\s*|\sae\s*|sae\s*:\s*|sae\s*nº.|sae\s*nº\s*+:|sae\s*nro:\s*|sae\s*+nro.|carta\s*:|carta\s*+:|SAE\s*N.\s*ª\s*|SAE\s*nª\s*|SAE\s*N.\s*º\s*|N°\s*)\s*(\d{8})', texto, re.IGNORECASE)
        resultado = re.search(
            r'\s*Resultado\s*+:\s*([^\n]+)', texto, re.IGNORECASE)

        datos.append({
            "titulo": ruta,
            "fecha": fecha.group(1).strip() if fecha else "S/D",
            "hora": hora.group(2).strip() if hora else "S/D",
            "causa": causa.group(1).strip() if causa else "S/D",
            "sae": sae.group(1).strip() if sae else "S/D",
            "jefe de servicio": jefe.group(1).strip() if jefe else "S/D",
            "oficial de servicio": oficial.group(1).strip() if oficial else "S/D",
            "operador de camara": operador.group(1).strip() if operador else "S/D",
            "resultado": resultado.group(1).strip() if resultado else "S/D"
        })

        # Extraer el bloque de reseña hasta "resultado" y quitar saltos de línea
        # Corregir la segunda expresión regular
        resena_bloque = re.search(
            r'(?:rese[nñ]a|rese[ñn]a)\s*:?\s*(.*?)Resultado\s*:',
            texto, re.IGNORECASE | re.DOTALL
        )
        if resena_bloque:
            resena_limpia = re.sub(r'\s+', ' ', resena_bloque.group(1)).strip()
        else:
            resena_limpia = "S/D"

        celda_combinada = (
            f"Fecha: {fecha.group(1).strip() if fecha else 'S/D'}\n"
            f"Reseña: {resena_limpia}\n"
            f"Resultado: {resultado.group(1).strip() if resultado else 'S/D'}"
        )

        datos_2.append(celda_combinada)

    df = pd.DataFrame(datos)
    df.to_excel("contenido_archivos.xlsx", index=False)

    # Guardar datos_2 en un archivo Word
    doc = Document()
    doc.add_heading('Reseñas extraídas', 0)
    for item in datos_2:
        doc.add_paragraph(item)
        doc.add_paragraph(
            '--------------------------------------------------------------')
    doc.save("reseña_archivos.docx")
    print("Datos extraídos y guardados en contenido_archivos.xlsx, reseña_archivos.docx")


# Initialize the TkinterDnD window
root = TkinterDnD.Tk()
root.title("Extraccion de datos .PDF a: -Parte semanal-Detenidos Oficiales-")
root.geometry("600x200")

# Create a frame to hold our content
frame = tk.Frame(root)
frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

# Create a label that can receive dropped files
ventana_pegado = tk.Label(
    frame,
    text="Copia y pega archivos aquí",
    bg="#2b74be",
    font=("Helvetica", 14),
    relief="ridge",
    height=4
)
ventana_pegado.pack(fill=tk.BOTH, expand=True)

# Add the "Aceptar" button
boton_aceptar = tk.Button(frame, text="Aceptar archivos", command=aceptar)
boton_aceptar.pack(side=tk.LEFT)

boton_cerrar = tk.Button(frame, text="Cerrar", command=exit)
boton_cerrar.pack(side=tk.RIGHT)

boton_ejecutar = tk.Button(frame, text="Ejecutar", command=extraer_convertir)
boton_ejecutar.pack(side=tk.BOTTOM)

# Register the label as a drop target
ventana_pegado.drop_target_register(DND_FILES)
ventana_pegado.dnd_bind('<<Drop>>', pegar)

# Run the application
if __name__ == "__main__":
    root.mainloop()
